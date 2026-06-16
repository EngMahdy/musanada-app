#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Professional Document Formatter
================================
تحسينات تصميم احترافية للفورمات:
- Header/Footer متقدم
- Page numbers
- Watermark
- Reference ID
- توقيع + ختم في صفحة واحدة
"""

from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_professional_header(doc, logo_path, company_name_ar, company_name_en, ref_id=None):
    """
    Header احترافي مع شعار + اسم الشركة + Reference ID
    """
    section = doc.sections[0]
    header = section.header
    
    # مسح المحتوى القديم
    for para in header.paragraphs:
        para.clear()
    
    # جدول 3 أعمدة: Logo | Company Name | Ref + Date
    table = header.add_table(rows=1, cols=3, width=Inches(7.5))
    table.autofit = False
    
    # العمود 1: Logo
    cell_logo = table.cell(0, 0)
    cell_logo.width = Inches(1.2)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if logo_path and Path(logo_path).exists():
        try:
            run_logo = p_logo.add_run()
            run_logo.add_picture(str(logo_path), width=Inches(0.9))
        except:
            p_logo.add_run("[LOGO]")
    
    # العمود 2: اسم الشركة
    cell_name = table.cell(0, 1)
    cell_name.width = Inches(4.5)
    p_name = cell_name.paragraphs[0]
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Arabic name (bold, larger)
    run_ar = p_name.add_run(company_name_ar + "\n")
    run_ar.font.size = Pt(14)
    run_ar.font.bold = True
    run_ar.font.color.rgb = RGBColor(184, 134, 11)  # ذهبي
    
    # English name (smaller)
    run_en = p_name.add_run(company_name_en)
    run_en.font.size = Pt(9)
    run_en.font.color.rgb = RGBColor(25, 25, 112)  # أزرق داكن
    
    # العمود 3: Reference + Date
    cell_ref = table.cell(0, 2)
    cell_ref.width = Inches(1.8)
    p_ref = cell_ref.paragraphs[0]
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    today = datetime.now().strftime("%d-%b-%Y")
    
    run_date = p_ref.add_run(f"Date: {today}\n")
    run_date.font.size = Pt(8)
    
    if ref_id:
        run_ref_id = p_ref.add_run(f"Ref: {ref_id}")
        run_ref_id.font.size = Pt(8)
        run_ref_id.font.bold = True
    
    # إضافة border سفلي للـheader
    set_cell_border(
        table.cell(0, 0),
        bottom={"sz": 12, "val": "single", "color": "B8860B"}
    )
    set_cell_border(
        table.cell(0, 1),
        bottom={"sz": 12, "val": "single", "color": "B8860B"}
    )
    set_cell_border(
        table.cell(0, 2),
        bottom={"sz": 12, "val": "single", "color": "B8860B"}
    )


def add_professional_footer(doc, company_data):
    """
    Footer احترافي مع معلومات الاتصال + رقم الصفحة
    """
    section = doc.sections[0]
    footer = section.footer
    
    # مسح القديم
    for para in footer.paragraphs:
        para.clear()
    
    # سطر معلومات الشركة
    p_info = footer.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_parts = []
    
    if company_data.get('legal_name'):
        contact_parts.append(company_data['legal_name'])
    
    if company_data.get('phone'):
        contact_parts.append(f"Tel: {company_data['phone']}")
    
    if company_data.get('email'):
        contact_parts.append(f"Email: {company_data['email']}")
    
    if company_data.get('license_number'):
        contact_parts.append(f"License: {company_data['license_number']}")
    
    info_text = " | ".join(contact_parts)
    
    run_info = p_info.add_run(info_text)
    run_info.font.size = Pt(7)
    run_info.font.color.rgb = RGBColor(128, 128, 128)
    
    # سطر رقم الصفحة
    p_page = footer.add_paragraph()
    p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_page = p_page.add_run("Page ")
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = RGBColor(128, 128, 128)
    
    # إضافة رقم الصفحة (XML)
    add_page_number_run(p_page)


def add_page_number_run(paragraph):
    """إضافة رقم الصفحة الديناميكي"""
    run = paragraph.add_run()
    
    # إنشاء field code لرقم الصفحة
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # تنسيق
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)


def add_watermark(doc, logo_path, opacity=0.1):
    """
    إضافة watermark شفاف من الشعار
    
    Note: docx-python لا يدعم watermarks مباشرة، هذه محاولة عبر header
    """
    # طريقة بديلة: إضافة الشعار كصورة شفافة في الخلفية
    # يحتاج معالجة XML متقدمة - نتجاوزها الآن
    pass


def add_signature_stamp_block(doc, signature_path, stamp_path, signatory_name, company_name, license_no):
    """
    كتلة توقيع + ختم احترافية في نهاية المستند
    """
    doc.add_paragraph()  # مسافة
    
    # جدول 2 أعمدة: التوقيع | الختم
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # خلية التوقيع
    cell_sig = table.cell(0, 0)
    p_sig = cell_sig.paragraphs[0]
    p_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if signature_path and Path(signature_path).exists():
        try:
            run_sig = p_sig.add_run()
            run_sig.add_picture(str(signature_path), width=Inches(2.0))
            p_sig.add_run("\n")
        except:
            pass
    
    # اسم الموقع
    run_name = p_sig.add_run(f"{signatory_name}\n")
    run_name.font.size = Pt(11)
    run_name.font.bold = True
    
    # اسم الشركة
    run_company = p_sig.add_run(f"{company_name}\n")
    run_company.font.size = Pt(9)
    
    # رقم الرخصة
    run_license = p_sig.add_run(f"License No: {license_no}")
    run_license.font.size = Pt(8)
    run_license.font.color.rgb = RGBColor(100, 100, 100)
    
    # خلية الختم
    cell_stamp = table.cell(0, 1)
    p_stamp = cell_stamp.paragraphs[0]
    p_stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if stamp_path and Path(stamp_path).exists():
        try:
            run_stamp = p_stamp.add_run()
            run_stamp.add_picture(str(stamp_path), width=Inches(1.8))
        except:
            pass


def set_cell_border(cell, **kwargs):
    """
    تعيين border لخلية جدول
    
    Args:
        cell: خلية الجدول
        **kwargs: top, bottom, left, right
                  كل واحد = {"sz": size, "val": type, "color": hex}
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            border_el = OxmlElement(f'w:{edge}')
            for key, value in kwargs[edge].items():
                border_el.set(qn(f'w:{key}'), str(value))
            tcBorders.append(border_el)
    
    tcPr.append(tcBorders)


def generate_reference_id(tender_no, year=None):
    """
    توليد Reference ID احترافي
    
    Format: MUS/BID/P236/2024
    """
    if year is None:
        year = datetime.now().year
    
    # تنظيف tender_no
    safe_tender = tender_no.replace(" ", "").replace("/", "-")[:20]
    
    return f"MUS/BID/{safe_tender}/{year}"


# ====== اختبار ======
if __name__ == "__main__":
    # اختبار Reference ID
    ref_id = generate_reference_id("P-236")
    print(f"Reference ID: {ref_id}")
    
    # اختبار إنشاء مستند
    doc = Document()
    
    company_data = {
        "legal_name": "مساندة للاستشارات الهندسية ودراسات الجدوى",
        "phone": "+971 2 123 4567",
        "email": "info@musanada.ae",
        "license_number": "CN-1234567"
    }
    
    add_professional_header(
        doc,
        logo_path=None,
        company_name_ar="مساندة للاستشارات الهندسية ودراسات الجدوى",
        company_name_en="MUSANADA Engineering Consultancy & Feasibility Studies",
        ref_id="MUS/BID/P236/2024"
    )
    
    add_professional_footer(doc, company_data)
    
    # محتوى تجريبي
    doc.add_heading("Form A — Proposal Cover Letter", 1)
    doc.add_paragraph("This is a test document with professional formatting.")
    
    # حفظ
    output_path = Path("/tmp/test_professional_format.docx")
    doc.save(str(output_path))
    print(f"✅ Test document saved: {output_path}")
