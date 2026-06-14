#!/usr/bin/env python3
"""
fill_dmt_forms.py — Fills ORIGINAL DMT auction forms with real bidder data.

Unlike generate_adio_forms.py (which builds new docx from scratch with placeholders),
this script TAKES the official DMT templates and FILLS the actual blank fields
with real values from bidder_data.json + tender_intelligence.json.

Forms handled:
- Form A: Letter of Auction (fills [Date], Auction Title, Bidder name, signature block)
- Form B: Experience & Capabilities (fills the projects table with real past projects)
- Form H: Non-Conflict of Interest Declaration (fills date, Auction Title, Bidder name)
- KYC Form: Investor Know Your Customer (fills 50+ company detail fields)
- NDU Form: Non-Disclosure Undertaking (fills Investor name)

Usage:
  python3 fill_dmt_forms.py <bidder_data.json> <tender_intelligence.json> <forms_source_dir> <output_dir>

Where:
- bidder_data.json: company info + projects (from form_snapshot)
- tender_intelligence.json: from tender_ai_reader.py
- forms_source_dir: directory with the 5 original DMT .docx files
- output_dir: where filled forms will be saved
"""

import sys
import json
import shutil
import re
from pathlib import Path
from datetime import datetime
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# =============== TEXT REPLACEMENT HELPERS ===============

def replace_text_in_paragraph(paragraph, replacements: dict):
    """Replace text in paragraph while preserving formatting.
    
    Handles cases where text spans multiple runs (common in Word documents).
    """
    # Strategy: collect all runs' text, do replacement on combined text,
    # then split back into runs preserving first run's formatting
    
    if not paragraph.runs:
        return False
    
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    changed = False
    
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, str(new))
            changed = True
    
    if changed:
        # Put all text in first run, clear others
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    
    return changed


def replace_in_document(doc, replacements: dict):
    """Walk through entire document (paragraphs + tables + textboxes) and apply replacements."""
    n = 0
    
    # Body paragraphs
    for p in doc.paragraphs:
        if replace_text_in_paragraph(p, replacements):
            n += 1
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_text_in_paragraph(p, replacements):
                        n += 1
    
    # Headers / footers
    for section in doc.sections:
        for hdr in [section.header, section.first_page_header, section.even_page_header]:
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                if replace_text_in_paragraph(p, replacements):
                    n += 1
        for ftr in [section.footer, section.first_page_footer, section.even_page_footer]:
            if ftr is None:
                continue
            for p in ftr.paragraphs:
                if replace_text_in_paragraph(p, replacements):
                    n += 1
    
    # Text inside text boxes (DMT KYC has these on cover page)
    n += _replace_in_textboxes(doc, replacements)
    
    return n


def _replace_in_textboxes(doc, replacements: dict):
    """Replace text inside <w:txbxContent> elements (Word text boxes)."""
    n = 0
    body = doc.element.body
    for tbc in body.iter(qn('w:txbxContent')):
        # Iterate over <w:t> elements inside the textbox
        text_elements = list(tbc.iter(qn('w:t')))
        if not text_elements:
            continue
        
        # Combine all text in the textbox to do replacement across runs
        combined = "".join(t.text or "" for t in text_elements)
        changed = False
        new_combined = combined
        for old, new in replacements.items():
            if old in new_combined:
                new_combined = new_combined.replace(old, str(new))
                changed = True
        
        if changed:
            # Put all text in first run, clear others
            text_elements[0].text = new_combined
            for t in text_elements[1:]:
                t.text = ""
            n += 1
    
    return n


# =============== INSERT TEXT AFTER A LABEL ===============

def insert_text_after_label(paragraph, label: str, value: str):
    """If paragraph text starts with `label` (e.g. 'Auction Title:'),
    append `value` after it (preserving the label)."""
    full_text = "".join(run.text for run in paragraph.runs)
    
    # Match the label with optional trailing colon and tabs/spaces
    pattern = re.compile(r'^(\s*' + re.escape(label.rstrip(":").strip()) + r'\s*:?\s*)(.*)$', re.IGNORECASE)
    m = pattern.match(full_text)
    
    if m:
        prefix = m.group(1)
        existing = m.group(2).strip()
        if not existing or existing in ["", "\t"]:
            # Empty - fill it
            new_text = prefix + str(value)
            if paragraph.runs:
                paragraph.runs[0].text = new_text
                for run in paragraph.runs[1:]:
                    run.text = ""
            return True
    return False


def fill_field_after_label(doc, label: str, value: str):
    """Search all paragraphs/cells for a label and fill after it."""
    filled = 0
    
    for p in doc.paragraphs:
        if insert_text_after_label(p, label, value):
            filled += 1
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if insert_text_after_label(p, label, value):
                        filled += 1
    
    return filled


# =============== IMAGE INSERTION ===============

def insert_image_in_signature_area(doc, image_path: str, marker_text: str, width_in_inches: float = 1.5):
    """Insert an image (signature/stamp) near a paragraph containing marker_text."""
    from docx.shared import Inches
    
    if not image_path or not Path(image_path).exists():
        return False
    
    # Search body paragraphs
    for p in doc.paragraphs:
        if marker_text.lower() in p.text.lower():
            try:
                # Add new paragraph after the marker with image
                run = p.add_run()
                run.add_picture(image_path, width=Inches(width_in_inches))
                return True
            except Exception as e:
                print(f"  Image insert error: {e}", file=sys.stderr)
                return False
    
    # Try in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if marker_text.lower() in p.text.lower():
                        try:
                            run = p.add_run()
                            run.add_picture(image_path, width=Inches(width_in_inches))
                            return True
                        except Exception:
                            return False
    return False


# =============== FORM A: LETTER OF AUCTION ===============

def fill_form_a(template_path: Path, output_path: Path, bidder: dict, intel: dict):
    """Fill Form A: Letter of Auction.
    
    Fields to fill:
    - [Date] → today's date
    - Auction Title: → from intel.authority.tender_title_en
    - [Bidder] → company legal name (appears in multiple places)
    - Signature/witness blocks
    """
    doc = Document(str(template_path))
    
    submission_date = bidder.get("submission_date") or datetime.now().strftime("%d %B %Y")
    bidder_name = bidder.get("company_legal_name", "[Bidder Name]")
    tender_title = (intel.get("authority", {}).get("tender_title_en")
                    or intel.get("project", {}).get("facility_type", "Tender"))
    signatory_name = bidder.get("authorized_signatory_name", "")
    signatory_title = bidder.get("authorized_signatory_title", "")
    
    # Replace text placeholders
    replacements = {
        "[Date]": submission_date,
        "[Bidder]": bidder_name,
        "[BIDDER]": bidder_name,
        "[Print name of Authorized Representative]": signatory_name,
        "[Position]": signatory_title,
    }
    n = replace_in_document(doc, replacements)
    print(f"  Form A: {n} replacements done")
    
    # Fill labeled fields
    filled = fill_field_after_label(doc, "Auction Title", tender_title)
    print(f"  Form A: filled {filled} 'Auction Title' field(s)")
    
    # Fill signature block table
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if "Print name of Authorized Representative" in cell_text:
                    # Add the name in a new paragraph
                    if signatory_name:
                        cell.add_paragraph(signatory_name)
                elif cell_text == "Position":
                    if signatory_title:
                        cell.add_paragraph(signatory_title)
    
    # Insert signature image near signature block
    sig_img = bidder.get("signature_image_path") or bidder.get("sig_path")
    if sig_img and Path(sig_img).exists():
        insert_image_in_signature_area(doc, sig_img,
            "Signature of Authorized Representative", width_in_inches=1.8)
    
    # Insert stamp
    stamp_img = bidder.get("stamp_image_path") or bidder.get("stamp_path")
    if stamp_img and Path(stamp_img).exists():
        insert_image_in_signature_area(doc, stamp_img, "Position", width_in_inches=1.2)
    
    doc.save(str(output_path))
    return n + filled


# =============== FORM B: EXPERIENCE & CAPABILITIES ===============

def fill_form_b(template_path: Path, output_path: Path, bidder: dict, intel: dict):
    """Fill Form B: Experience & Capabilities.
    
    The form has a 7-column projects table:
    No. | Name & location | Client name/contact | Bidder scope | Value AED | Start | Completion
    
    We fill rows 1-3 (and extend if more projects provided).
    """
    doc = Document(str(template_path))
    
    projects = bidder.get("projects_completed", []) or []
    
    if not doc.tables:
        print("  Form B: no table found!", file=sys.stderr)
        doc.save(str(output_path))
        return 0
    
    # Find the main projects table (8 rows × 7 cols)
    main_table = None
    for t in doc.tables:
        if len(t.columns) >= 6:
            main_table = t
            break
    
    if not main_table:
        doc.save(str(output_path))
        return 0
    
    # Fill rows 1-3 (after header row 0) with first 3 projects
    filled_rows = 0
    for i, project in enumerate(projects[:3]):
        if i + 1 >= len(main_table.rows):
            break
        
        row = main_table.rows[i + 1]
        cells = row.cells
        
        if len(cells) >= 7:
            cells[0].text = str(i + 1)
            cells[1].text = f"{project.get('name', '')}\n{project.get('location', '')}"
            cells[2].text = f"{project.get('client', '')}\n{project.get('client_contact', '')}"
            cells[3].text = project.get('scope', '')
            
            amt = project.get('amount', 0) or project.get('amount_aed', 0)
            if isinstance(amt, (int, float)) and amt > 0:
                cells[4].text = f"AED {int(amt):,}"
            
            cells[5].text = str(project.get('start', '') or project.get('start_date', ''))
            cells[6].text = str(project.get('end', '') or project.get('completion_date', '') or project.get('end_date', ''))
            
            filled_rows += 1
    
    print(f"  Form B: filled {filled_rows}/3 project rows (of {len(projects)} total provided)")
    
    # Add extra rows if more projects (extend table)
    if len(projects) > 3:
        # Find the "Sub-Consulted" header row to know where to insert
        for extra in projects[3:10]:  # max 10 total
            new_row = main_table.add_row()
            cells = new_row.cells
            if len(cells) >= 7:
                cells[0].text = str(filled_rows + 1)
                cells[1].text = f"{extra.get('name', '')}\n{extra.get('location', '')}"
                cells[2].text = f"{extra.get('client', '')}\n{extra.get('client_contact', '')}"
                cells[3].text = extra.get('scope', '')
                amt = extra.get('amount', 0) or extra.get('amount_aed', 0)
                if isinstance(amt, (int, float)) and amt > 0:
                    cells[4].text = f"AED {int(amt):,}"
                cells[5].text = str(extra.get('start', '') or extra.get('start_date', ''))
                cells[6].text = str(extra.get('end', '') or extra.get('completion_date', ''))
                filled_rows += 1
    
    doc.save(str(output_path))
    return filled_rows


# =============== FORM H: NON-CONFLICT DECLARATION ===============

def fill_form_h(template_path: Path, output_path: Path, bidder: dict, intel: dict):
    """Fill Form H: Non-Conflict of Interest Declaration."""
    doc = Document(str(template_path))
    
    submission_date = bidder.get("submission_date") or datetime.now().strftime("%d %B %Y")
    bidder_name = bidder.get("company_legal_name", "[Bidder Name]")
    tender_title = (intel.get("authority", {}).get("tender_title_en")
                    or intel.get("project", {}).get("facility_type", "Tender"))
    signatory_name = bidder.get("authorized_signatory_name", "")
    signatory_title = bidder.get("authorized_signatory_title", "")
    
    # The form uses "Bidder" placeholders in various forms
    # Note: the original has "…….Bidder ………." which we need to replace
    replacements = {
        "…….Bidder ……….": bidder_name,
        "..…..Bidder……….": bidder_name,
        "[Bidder]": bidder_name,
        "..........": bidder_name,  # generic blanks
    }
    n = replace_in_document(doc, replacements)
    
    # Fill labeled fields
    filled = 0
    filled += fill_field_after_label(doc, "Auction Title", tender_title)
    
    # Fill By: and Title:
    for p in doc.paragraphs:
        text = p.text
        if text.startswith("By:") and "__" in text:
            new_text = f"By: {signatory_name}"
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
            filled += 1
        elif text.startswith("Title:") and "__" in text:
            new_text = f"Title: {signatory_title}"
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
            filled += 1
    
    # Insert signature near By:
    sig_img = bidder.get("signature_image_path") or bidder.get("sig_path")
    if sig_img and Path(sig_img).exists():
        insert_image_in_signature_area(doc, sig_img, "By:", width_in_inches=1.8)
    
    # Insert stamp near bottom
    stamp_img = bidder.get("stamp_image_path") or bidder.get("stamp_path")
    if stamp_img and Path(stamp_img).exists():
        insert_image_in_signature_area(doc, stamp_img, "Title:", width_in_inches=1.2)
    
    print(f"  Form H: {n} replacements + {filled} labeled fields")
    
    doc.save(str(output_path))
    return n + filled


# =============== KYC FORM ===============

def fill_kyc(template_path: Path, output_path: Path, bidder: dict, intel: dict):
    """Fill the KYC form's many labeled fields."""
    doc = Document(str(template_path))
    
    submission_date = bidder.get("submission_date") or datetime.now().strftime("%d %B %Y")
    
    # First do simple text replacements for placeholders
    simple_replacements = {
        "[Date]": submission_date,
    }
    n_replace = replace_in_document(doc, simple_replacements)
    
    # Build a map of all KYC field labels → values
    field_map = {
        # Company Details (Section A)
        "Full legal name of company": bidder.get("company_legal_name", ""),
        "Company license number": bidder.get("trade_license_no", ""),
        "Country of incorporation": "United Arab Emirates",
        "Date of incorporation": bidder.get("establishment_date", ""),
        "Legal Form: (LLC, PJSC, etc.)": bidder.get("legal_form", ""),
        "Legal Form": bidder.get("legal_form", ""),
        "Tax Residency (Country/Countries)": "United Arab Emirates",
        "Tax Residency": "United Arab Emirates",
        "Principal place of business": bidder.get("hq_address", ""),
        "Overview of your business operations, including key products or services offered.": bidder.get("nature_of_business", ""),
        "Economic Licence No.": bidder.get("trade_license_no", ""),
        "VAT registration No.": bidder.get("vat_no", ""),
        "Date of Establishment / D.O.B": bidder.get("establishment_date", ""),
        "Registered Business Address of Headquarters": bidder.get("hq_address", ""),
    }
    
    # Fill labeled fields
    filled_count = n_replace
    for label, value in field_map.items():
        if value:
            n = fill_field_after_label(doc, label, value)
            filled_count += n
    
    # Section B Primary Contact - appears as labeled list "Name:" "Title:" "Phone:" "Email:"
    # We need to find them inside the "Primary Contact Person" section
    primary_contact = {
        "name": bidder.get("contact_person_name") or bidder.get("authorized_signatory_name", ""),
        "title": bidder.get("authorized_signatory_title", ""),
        "phone": bidder.get("contact_person_phone") or bidder.get("company_phone", ""),
        "email": bidder.get("contact_person_email") or bidder.get("company_email", ""),
    }
    authorized = {
        "name": bidder.get("authorized_signatory_name", ""),
        "title": bidder.get("authorized_signatory_title", ""),
        "phone": bidder.get("company_phone", ""),
        "email": bidder.get("company_email", ""),
    }
    
    # Walk through paragraphs and fill contact fields under their context
    section = ""
    for p in doc.paragraphs:
        text = p.text.strip()
        
        if "Primary Contact Person" in text:
            section = "primary"
        elif "Authorized Signatory" in text and ":" in text:
            section = "authorized"
        elif "advisor" in text.lower() and "name of advisor" in text.lower():
            section = "advisor"
        
        # Now fill fields based on current section
        if section == "primary":
            for label, key in [("Name", "name"), ("Title", "title"), ("Phone", "phone"), ("Email", "email")]:
                if text.startswith(f"{label}:"):
                    val = primary_contact.get(key, "")
                    if val and (text == f"{label}:" or text.endswith(f"{label}:")):
                        if p.runs:
                            p.runs[0].text = f"{label}: {val}"
                            for r in p.runs[1:]: r.text = ""
                            filled_count += 1
                            break
        elif section == "authorized":
            for label, key in [("Name", "name"), ("Title", "title"), ("Phone", "phone"), ("Email", "email")]:
                if text.startswith(f"{label}:"):
                    val = authorized.get(key, "")
                    if val and (text == f"{label}:" or text.endswith(f"{label}:")):
                        if p.runs:
                            p.runs[0].text = f"{label}: {val}"
                            for r in p.runs[1:]: r.text = ""
                            filled_count += 1
                            break
    
    print(f"  KYC: filled {filled_count} fields")
    
    doc.save(str(output_path))
    return filled_count


# =============== NDU FORM ===============

def fill_ndu(template_path: Path, output_path: Path, bidder: dict, intel: dict):
    """Fill the NDU (Non-Disclosure Undertaking) — mostly inserting investor name."""
    doc = Document(str(template_path))
    
    bidder_name = bidder.get("company_legal_name", "[Investor Name]")
    project_name = (intel.get("authority", {}).get("tender_title_en")
                    or intel.get("project", {}).get("facility_type", "the Project"))
    
    replacements = {
        "[Investor]": bidder_name,
        "…………………………": bidder_name,
        "…………………": project_name,  # for "Project means the provision …"
    }
    n = replace_in_document(doc, replacements)
    
    print(f"  NDU: {n} replacements")
    
    doc.save(str(output_path))
    return n


# =============== MAIN ===============

def main():
    if len(sys.argv) != 5:
        print("Usage: fill_dmt_forms.py <bidder_data.json> <tender_intelligence.json> <forms_source_dir> <output_dir>")
        sys.exit(1)
    
    bidder_data_path = Path(sys.argv[1])
    intel_path = Path(sys.argv[2])
    source_dir = Path(sys.argv[3])
    output_dir = Path(sys.argv[4])
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    bidder = json.loads(bidder_data_path.read_text(encoding="utf-8"))
    intel = json.loads(intel_path.read_text(encoding="utf-8"))
    
    print(f"\n╔══ Musanada DMT Forms Filler ══╗")
    print(f"║ Bidder: {bidder.get('company_legal_name', '?')[:40]}")
    print(f"║ Tender: {intel.get('authority', {}).get('tender_title_en', '?')[:40]}")
    print(f"╚════════════════════════════════╝\n")
    
    # Map: form name → (source filename pattern, target filename, fill function)
    forms = [
        ("Form A", "Form A", "Form_A_Letter_of_Auction.docx", fill_form_a),
        ("Form B", "Form B", "Form_B_Experience_Capabilities.docx", fill_form_b),
        ("Form H", "Form H", "Form_H_Non_Conflict_Declaration.docx", fill_form_h),
        ("KYC", "KYC", "Investor_KYC_Form.docx", fill_kyc),
        ("NDU", "NDU", "Investor_NDU.docx", fill_ndu),
    ]
    
    # Find source files
    source_files = list(source_dir.glob("*.docx"))
    
    successful = 0
    for label, pattern, target_name, fill_fn in forms:
        # Find matching source file
        matched = None
        for sf in source_files:
            if pattern.lower() in sf.name.lower():
                matched = sf
                break
        
        if not matched:
            print(f"⚠ {label}: source not found (pattern: '{pattern}')")
            continue
        
        target_path = output_dir / target_name
        print(f"\n📄 Filling {label} from {matched.name}")
        try:
            result = fill_fn(matched, target_path, bidder, intel)
            size_kb = target_path.stat().st_size // 1024
            print(f"  ✓ Saved: {target_name} ({size_kb} KB)")
            successful += 1
        except Exception as e:
            import traceback
            print(f"  ✗ Error: {e}")
            traceback.print_exc()
    
    print(f"\n✓ Done: {successful}/{len(forms)} forms filled successfully")


if __name__ == "__main__":
    main()
