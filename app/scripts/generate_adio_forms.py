#!/usr/bin/env python3
"""
generate_adio_forms.py — يولد نماذج ADIO احترافية بصيغة DOCX
مع تنسيق أنيق (header, جداول، توقيع، ختم) من ملف JSON.

الاستخدام:
  python3 generate_adio_forms.py <bidder_data.json> <output_dir>

bidder_data.json يحتوي على بيانات المزايد + تفاصيل المناقصة.
"""

import sys
import json
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION, WD_ORIENTATION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: pip install python-docx")
    sys.exit(1)


# ===== TYPOGRAPHY & COLORS =====
COLOR_PRIMARY = RGBColor(0x0B, 0x3D, 0x7A)      # أزرق غامق (Royal Navy)
COLOR_ACCENT = RGBColor(0xC9, 0xA2, 0x4C)        # ذهبي (Gold)
COLOR_TEXT = RGBColor(0x1A, 0x1A, 0x1A)          # أسود ناعم
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)          # رمادي
COLOR_LIGHT = RGBColor(0xF5, 0xF5, 0xF5)         # رمادي فاتح (للـrows)


def set_cell_background(cell, color_hex):
    """Set cell background color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="0B3D7A", size=4):
    """Set thin borders on cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_styled_paragraph(doc, text, size=11, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def add_header_bar(doc, title, subtitle="", logo_path=None):
    """Add an elegant header bar with optional company logo."""
    # If logo provided, add a 2-cell table: logo | title
    if logo_path and Path(logo_path).exists():
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Logo cell
        logo_cell = table.rows[0].cells[0]
        logo_cell.width = Cm(4)
        set_cell_background(logo_cell, "FFFFFF")
        logo_p = logo_cell.paragraphs[0]
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            logo_run = logo_p.add_run()
            logo_run.add_picture(str(logo_path), width=Cm(3))
        except Exception as e:
            logo_p.add_run("[LOGO]").font.size = Pt(12)
        
        # Title cell
        title_cell = table.rows[0].cells[1]
        title_cell.width = Cm(13)
        set_cell_background(title_cell, "0B3D7A")
        p = title_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        
        if subtitle:
            sub_p = title_cell.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run(subtitle)
            sub_run.font.name = 'Calibri'
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(0xC9, 0xA2, 0x4C)
            sub_run.italic = True
        
        title_cell.paragraphs[0].paragraph_format.space_before = Pt(12)
        if subtitle:
            title_cell.paragraphs[1].paragraph_format.space_after = Pt(8)
    else:
        # Original single-cell header
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        set_cell_background(cell, "0B3D7A")
        
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.name = 'Calibri'
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        
        if subtitle:
            sub_p = cell.add_paragraph()
            sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub_p.add_run(subtitle)
            sub_run.font.name = 'Calibri'
            sub_run.font.size = Pt(11)
            sub_run.font.color.rgb = RGBColor(0xC9, 0xA2, 0x4C)
            sub_run.italic = True
        
        cell.paragraphs[0].paragraph_format.space_before = Pt(8)
        if subtitle:
            cell.paragraphs[1].paragraph_format.space_after = Pt(8)
        else:
            cell.paragraphs[0].paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph()  # spacer


def add_recipient_block(doc):
    """Add the ADIO recipient block."""
    block = doc.add_table(rows=5, cols=1)
    block.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    rows_data = [
        ("To: Head of Procurement", True),
        ("Abu Dhabi Investment Office (ADIO)", True),
        ("Abu Dhabi, United Arab Emirates", False),
        ("Attention: Procurement Department ADIO", False),
        ("Email: procurement.musataha@adio.gov.ae", False),
    ]
    for i, (txt, bold) in enumerate(rows_data):
        cell = block.rows[i].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.size = Pt(10)
        run.bold = bold
        if bold:
            run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph()


def add_meta_block(doc, data):
    """Add tender meta info block."""
    add_styled_paragraph(doc, "PROPOSAL DETAILS", size=12, bold=True, color=COLOR_PRIMARY, space_after=4)
    
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    rows = [
        ("Name of Bidder", data.get("company_legal_name", "________________")),
        ("Tender No.", data.get("tender_no", "________________")),
        ("Tender Name", data.get("tender_name", "________________")),
        ("Date", data.get("submission_date", datetime.now().strftime("%d-%m-%Y"))),
        ("Subject", data.get("subject", "________________")),
    ]
    
    for i, (k, v) in enumerate(rows):
        c1 = table.rows[i].cells[0]
        c2 = table.rows[i].cells[1]
        
        c1.width = Cm(5)
        c2.width = Cm(11)
        
        set_cell_background(c1, "0B3D7A")
        set_cell_borders(c1)
        set_cell_borders(c2)
        
        p1 = c1.paragraphs[0]
        run1 = p1.add_run(k)
        run1.font.size = Pt(10)
        run1.bold = True
        run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
        p2 = c2.paragraphs[0]
        run2 = p2.add_run(str(v))
        run2.font.size = Pt(10)
        run2.font.color.rgb = COLOR_TEXT
    
    doc.add_paragraph()


def add_section_header(doc, text):
    """Add a section header with gold accent line."""
    # gold line table
    line = doc.add_table(rows=1, cols=1)
    cell = line.rows[0].cells[0]
    set_cell_background(cell, "C9A24C")
    cell.height = Cm(0.1)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(1)
    
    add_styled_paragraph(doc, text, size=14, bold=True, color=COLOR_PRIMARY, space_after=8)


def add_signature_block(doc, data):
    """Add signature & stamp block at end. Supports inserted signature/stamp images."""
    doc.add_paragraph()
    add_section_header(doc, "AUTHORIZED SIGNATURE")
    
    addr_p = doc.add_paragraph()
    addr_p.add_run("Bidder's Business Address: ").bold = True
    addr_p.add_run(data.get("bidder_address", "____________________________"))
    
    doc.add_paragraph()
    
    # Use a 2-col layout: signature info | signature/stamp images
    sig_path = data.get("signature_image_path")
    stamp_path = data.get("stamp_image_path")
    has_images = (sig_path and Path(sig_path).exists()) or (stamp_path and Path(stamp_path).exists())
    
    if has_images:
        outer = doc.add_table(rows=1, cols=2)
        outer.alignment = WD_TABLE_ALIGNMENT.LEFT
        left_cell = outer.rows[0].cells[0]
        right_cell = outer.rows[0].cells[1]
        left_cell.width = Cm(10)
        right_cell.width = Cm(8)
        
        # Left: info table
        info_table = left_cell.add_table(rows=5, cols=2)
        rows = [
            ("Signature:", ""),
            ("Name:", data.get("authorized_signatory_name", "____________________")),
            ("Position:", data.get("authorized_signatory_title", "____________________")),
            ("Name of Company:", data.get("company_legal_name", "____________________")),
            ("Date:", data.get("submission_date", datetime.now().strftime("%d-%m-%Y"))),
        ]
        for i, (k, v) in enumerate(rows):
            c1 = info_table.rows[i].cells[0]
            c2 = info_table.rows[i].cells[1]
            c1.width = Cm(3.5)
            c2.width = Cm(6.5)
            run1 = c1.paragraphs[0].add_run(k)
            run1.bold = True
            run1.font.size = Pt(10)
            run2 = c2.paragraphs[0].add_run(v)
            run2.font.size = Pt(10)
        
        # Right: signature and stamp images
        if sig_path and Path(sig_path).exists():
            sig_p = right_cell.paragraphs[0]
            sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                sig_p.add_run().add_picture(str(sig_path), width=Cm(5))
            except:
                sig_p.add_run("[Signature]").italic = True
        
        if stamp_path and Path(stamp_path).exists():
            stamp_p = right_cell.add_paragraph()
            stamp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                stamp_p.add_run().add_picture(str(stamp_path), width=Cm(4))
            except:
                stamp_p.add_run("[Stamp]").italic = True
    else:
        sig_table = doc.add_table(rows=5, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        rows = [
            ("Signature:", "________________________________"),
            ("Name:", data.get("authorized_signatory_name", "________________________________")),
            ("Position:", data.get("authorized_signatory_title", "________________________________")),
            ("Name of Company:", data.get("company_legal_name", "________________________________")),
            ("Date:", data.get("submission_date", datetime.now().strftime("%d-%m-%Y"))),
        ]
        
        for i, (k, v) in enumerate(rows):
            c1 = sig_table.rows[i].cells[0]
            c2 = sig_table.rows[i].cells[1]
            c1.width = Cm(4)
            c2.width = Cm(12)
            
            run1 = c1.paragraphs[0].add_run(k)
            run1.bold = True
            run1.font.size = Pt(10)
            
            run2 = c2.paragraphs[0].add_run(v)
            run2.font.size = Pt(10)
        
        doc.add_paragraph()
        
        stamp = doc.add_paragraph()
        stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = stamp.add_run("[ COMPANY STAMP ]")
        run.italic = True
        run.font.color.rgb = COLOR_GRAY
        run.font.size = Pt(10)


def setup_page(doc, landscape=False):
    """Configure A4 page with proper margins to prevent print issues."""
    for section in doc.sections:
        if landscape:
            section.orientation = WD_ORIENTATION.LANDSCAPE
            section.page_width = Mm(297)
            section.page_height = Mm(210)
        else:
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)


def add_page_footer(doc, company_name=""):
    """Add a professional footer to every page."""
    section = doc.sections[0]
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"{company_name} | Confidential")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GRAY
    run.italic = True


# ============================================================
# FORM A — Proposal Cover Letter
# ============================================================
def build_form_a(data, output_dir):
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    
    add_header_bar(doc, "FORM A", "Proposal Cover Letter", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": "Form A — Proposal Cover Letter"}
    add_meta_block(doc, meta)

    # ===== Pricing offer paragraph (auto-filled from Tender Intelligence) =====
    _proposed = data.get("proposed_annual_rent_aed")
    _per_sqm = data.get("proposed_rent_per_sqm")
    _src = data.get("rent_source", "")
    _area = data.get("plot_area_sqm")
    _lease = data.get("lease_years", 25)
    _irr = data.get("irr_pct")
    _payback = data.get("payback_years")
    if _proposed and _per_sqm and _area:
        add_section_header(doc, "PROPOSED FINANCIAL OFFER")
        offer_text = (
            f"In line with our comprehensive feasibility analysis, we hereby submit our "
            f"financial offer for the captioned auction on the basis of the following terms:\n\n"
            f"• Proposed Annual Rent to Authority: AED {_proposed:,.0f}\n"
            f"• Equivalent Rate: AED {_per_sqm:,.0f} per sqm per year\n"
            f"• Plot Area: {_area:,.0f} sqm\n"
            f"• Lease Term: {_lease} years\n"
        )
        if _irr is not None:
            offer_text += f"• Projected Project IRR: {_irr}%\n"
        if _payback is not None:
            offer_text += f"• Projected Payback Period: {_payback} years\n"
        offer_text += (
            f"\nThis pricing has been derived through a structured financial model "
            f"based on {_src.lower() if _src else 'market research'}, and is supported by the "
            f"detailed Feasibility Study (Volume III) submitted herewith."
        )
        for para in offer_text.split("\n"):
            if para.strip():
                add_styled_paragraph(doc, para.strip(), size=10, space_after=4)

    # CAPEX section
    add_section_header(doc, "CAPITAL EXPENDITURES BREAKDOWN")
    
    capex = data.get("capex", {})
    
    # Helper to add CAPEX subsection
    def add_capex_table(title, items_dict, total_key=None):
        add_styled_paragraph(doc, title, size=11, bold=True, color=COLOR_PRIMARY, space_after=4)
        
        items = [(k, v) for k, v in items_dict.items() if k != "total"]
        if total_key:
            total = items_dict.get(total_key, sum(v for k, v in items if isinstance(v, (int, float))))
        else:
            total = items_dict.get("total", sum(v for k, v in items if isinstance(v, (int, float))))
        
        table = doc.add_table(rows=len(items) + 1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        h1 = table.rows[0].cells[0]
        h2 = table.rows[0].cells[1]
        set_cell_background(h1, "0B3D7A")
        set_cell_background(h2, "0B3D7A")
        set_cell_borders(h1)
        set_cell_borders(h2)
        run1 = h1.paragraphs[0].add_run("Item")
        run1.bold = True
        run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run1.font.size = Pt(10)
        h2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = h2.paragraphs[0].add_run("Total Cost (AED)")
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(10)
        
        # Data rows
        for i, (k, v) in enumerate(items, start=1):
            c1 = table.rows[i].cells[0]
            c2 = table.rows[i].cells[1]
            set_cell_borders(c1)
            set_cell_borders(c2)
            
            if i % 2 == 0:
                set_cell_background(c1, "F5F5F5")
                set_cell_background(c2, "F5F5F5")
            
            c1.paragraphs[0].add_run(k).font.size = Pt(10)
            c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if isinstance(v, (int, float)):
                v_str = f"AED {v:,}"
            else:
                v_str = str(v)
            c2.paragraphs[0].add_run(v_str).font.size = Pt(10)
        
        # Total row (bold, accent)
        total_row = table.add_row()
        tc1 = total_row.cells[0]
        tc2 = total_row.cells[1]
        set_cell_background(tc1, "C9A24C")
        set_cell_background(tc2, "C9A24C")
        set_cell_borders(tc1)
        set_cell_borders(tc2)
        run1 = tc1.paragraphs[0].add_run(f"Total {title.split(chr(46))[-1].strip()}")
        run1.bold = True
        run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run1.font.size = Pt(10)
        tc2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if isinstance(total, (int, float)):
            t_str = f"AED {total:,}"
        else:
            t_str = str(total)
        run2 = tc2.paragraphs[0].add_run(t_str)
        run2.bold = True
        run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run2.font.size = Pt(10)
        
        doc.add_paragraph()
        return total if isinstance(total, (int, float)) else 0
    
    grand = 0
    grand += add_capex_table("A. Preparation", capex.get("preparation", {"Preliminary studies": "_____", "Site preparation": "_____"}))
    grand += add_capex_table("B. Construction", capex.get("construction", {"Shops construction": "_____", "Green area": "_____", "Restroom": "_____", "Parking and roads": "_____"}))
    grand += add_capex_table("C. Project Management", capex.get("project_management", {"Project management": "_____", "Pre-opening expenses": "_____", "Government fees": "_____"}))
    grand += add_capex_table("D. Information Technology", capex.get("it", {"Networking and connectivity": "_____"}))
    
    # Grand total
    add_section_header(doc, "TOTAL CAPITAL EXPENDITURES")
    
    grand_table = doc.add_table(rows=1, cols=2)
    gc1 = grand_table.rows[0].cells[0]
    gc2 = grand_table.rows[0].cells[1]
    set_cell_background(gc1, "0B3D7A")
    set_cell_background(gc2, "0B3D7A")
    set_cell_borders(gc1, color="C9A24C", size=8)
    set_cell_borders(gc2, color="C9A24C", size=8)
    run1 = gc1.paragraphs[0].add_run("TOTAL CAPITAL EXPENDITURES")
    run1.bold = True
    run1.font.size = Pt(13)
    run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    gc2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_str = f"AED {grand:,}" if grand else f"AED {capex.get('total', '_____')}"
    run2 = gc2.paragraphs[0].add_run(total_str)
    run2.bold = True
    run2.font.size = Pt(13)
    run2.font.color.rgb = RGBColor(0xC9, 0xA2, 0x4C)
    
    doc.add_paragraph()
    
    # Declaration
    add_section_header(doc, "BIDDER DECLARATION")
    declaration_text = (
        f"We, {data.get('company_legal_name', '________________')}, hereby submit this proposal in response "
        f"to the Request for Proposal (RFP) referenced above. We confirm that:\n\n"
        "1. We have reviewed all RFP documents and accept all terms and conditions.\n"
        "2. This proposal is binding upon us for a period of 120 days from the submission date.\n"
        "3. All capital expenditure figures provided are accurate to the best of our knowledge.\n"
        "4. We commit to delivering the project in accordance with the timeline and scope specified."
    )
    for line in declaration_text.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(4)
            for run in p.runs:
                run.font.size = Pt(10)
    
    add_signature_block(doc, data)
    
    out = Path(output_dir) / "Form_A_Cover_Letter.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# FORM D — Bidder's Details
# ============================================================
def build_form_d(data, output_dir):
    doc = Document()
    setup_page(doc)
    add_page_footer(doc, data.get('company_legal_name', ''))
    
    add_header_bar(doc, "FORM D", "Bidder's Details and Experience", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": "Form D — Bidder's Details"}
    add_meta_block(doc, meta)
    
    add_section_header(doc, "BIDDER'S DETAILS")
    
    # Section 1 (items 1-12)
    rows_data = [
        ("1", "Company full name", data.get("company_legal_name", "")),
        ("2", "Company short name (if any)", data.get("company_short_name", "")),
        ("3", "Legal form", data.get("legal_form", "")),
        ("4", "Nature of business", data.get("nature_of_business", "")),
        ("5", "Head office address and P.O. Box number", data.get("hq_address", "")),
        ("6", "Nationality of partners", data.get("partners_nationality", "")),
        ("7", "Date of establishment", data.get("establishment_date", "")),
        ("8", "Certificate of incorporation and Commercial License with related activities and services in the emirate of Abu Dhabi (provide copy)", "COPY ATTACHED"),
        ("9", "Power of Attorney holder (provide copy)", "COPY ATTACHED"),
        ("10", "Telephone number", data.get("company_phone", "")),
        ("11", "Facsimile number", data.get("company_fax", "")),
        ("12", "Email address", data.get("company_email", "")),
        ("13", "Web access", data.get("company_website", "")),
        ("14", "Named contact person", data.get("contact_person_name", "")),
        ("15", "Contact person telephone number", data.get("contact_person_phone", "")),
        ("16", "Contact person email address", data.get("contact_person_email", "")),
        ("17", "Bidder's organization chart", "COPY ATTACHED"),
        ("18", f"Bidder's Experience (evidence to demonstrate minimum {data.get('min_exp_years', 2)} years of experience in {data.get('experience_domain', 'building and operating similar facilities')})", "COPY ATTACHED"),
    ]
    
    table = doc.add_table(rows=len(rows_data) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    headers = ["No.", "Requirement", "Detail"]
    widths = [Cm(1.5), Cm(8), Cm(7.5)]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        set_cell_background(c, "0B3D7A")
        set_cell_borders(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    
    # Data
    for i, (n, req, det) in enumerate(rows_data, start=1):
        for j, val in enumerate([n, req, det]):
            c = table.rows[i].cells[j]
            c.width = widths[j]
            set_cell_borders(c)
            if i % 2 == 0:
                set_cell_background(c, "F5F5F5")
            
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            if val == "COPY ATTACHED":
                run.bold = True
                run.font.color.rgb = COLOR_ACCENT
    
    doc.add_paragraph()
    
    # Attachments checklist
    add_section_header(doc, "REQUIRED ATTACHMENTS")
    
    att_data = [
        ("N8", "Certificate of Incorporation + Commercial License (with relevant activities)"),
        ("N9", "Power of Attorney for authorized signatory"),
        ("N17", "Organization Chart (visual structure with names & titles)"),
        ("N18", "Bidder's Experience Evidence (project list with proof)"),
        ("Bonus", "Detailed Company Profile / Brochure"),
    ]
    
    att_table = doc.add_table(rows=len(att_data) + 1, cols=3)
    att_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(["Item No.", "Attachment", "Status"]):
        c = att_table.rows[0].cells[i]
        set_cell_background(c, "0B3D7A")
        set_cell_borders(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    
    for i, (n, item) in enumerate(att_data, start=1):
        for j, val in enumerate([n, item, "☐ Attached"]):
            c = att_table.rows[i].cells[j]
            set_cell_borders(c)
            if i % 2 == 0:
                set_cell_background(c, "F5F5F5")
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    
    add_signature_block(doc, data)
    
    out = Path(output_dir) / "Form_D_Bidder_Details.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# FORM E — Sworn Declaration
# ============================================================
def build_form_e(data, output_dir):
    doc = Document()
    setup_page(doc)
    add_page_footer(doc, data.get('company_legal_name', ''))
    
    add_header_bar(doc, "FORM E", "Sworn Declaration", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": "Form E — Sworn Declaration"}
    add_meta_block(doc, meta)
    
    add_section_header(doc, "DECLARATION")
    
    intro = doc.add_paragraph()
    intro.add_run("I, ").font.size = Pt(11)
    name_run = intro.add_run(data.get("authorized_signatory_name", "________________"))
    name_run.bold = True
    name_run.font.size = Pt(11)
    name_run.font.color.rgb = COLOR_PRIMARY
    intro.add_run(", the undersigned, a duly authorized representative of ").font.size = Pt(11)
    company_run = intro.add_run(data.get("company_legal_name", "________________"))
    company_run.bold = True
    company_run.font.size = Pt(11)
    company_run.font.color.rgb = COLOR_PRIMARY
    intro.add_run(", do hereby depose and state under penalty of perjury that:").font.size = Pt(11)
    
    doc.add_paragraph()
    
    declarations = [
        ("1. Accuracy of Information",
         "All information and documentation submitted as part of this proposal — including but not limited to corporate details, financial statements, experience records, and supporting evidence — is true, accurate, and complete to the best of my knowledge and belief."),
        ("2. No Conflict of Interest",
         f"There does not exist any conflict of interest which would prevent {data.get('company_legal_name', 'the Bidder')}, or any of its shareholders, partners, directors, or affiliates, from entering into a contract with the Abu Dhabi Investment Office (ADIO) or with the Procuring Entity referenced in this RFP."),
        ("3. No Bankruptcy Proceedings",
         f"Neither {data.get('company_legal_name', 'the Bidder')}, nor any of its partners or shareholders, has been the subject of any bankruptcy proceedings, reorganization for the benefit of creditors, receivership, or similar insolvency proceedings within the last five (5) years preceding the date of this declaration."),
        ("4. No Material Adverse Changes",
         f"Since the date of the last audited reports and financial statements of {data.get('company_legal_name', 'the Bidder')}, or those of its shareholders or partners — submitted as part of this Proposal — there have been no material adverse changes that would affect the financial standing or operational capacity of the Bidder."),
        ("5. No Contract Disputes or Legal Proceedings",
         f"No contract disputes, arbitral proceedings, or material legal proceedings involving {data.get('company_legal_name', 'the Bidder')} have occurred in the last five (5) years preceding the date of this declaration."),
    ]
    
    for title, body in declarations:
        add_styled_paragraph(doc, title, size=11, bold=True, color=COLOR_PRIMARY, space_after=4)
        p = doc.add_paragraph(body)
        p.paragraph_format.space_after = Pt(10)
        for r in p.runs:
            r.font.size = Pt(10)
    
    doc.add_paragraph()
    add_section_header(doc, "CERTIFICATION")
    
    cert = doc.add_paragraph(
        f"I declare, certify, verify, and state under penalty of perjury that the foregoing is true and correct, "
        f"and that I am duly authorized to submit this letter and Proposal on behalf of {data.get('company_legal_name', '________________')}."
    )
    for r in cert.runs:
        r.font.size = Pt(10)
    cert.paragraph_format.space_after = Pt(20)
    
    add_signature_block(doc, data)
    
    out = Path(output_dir) / "Form_E_Sworn_Declaration.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# FORM I — Proposal Checklist (3 versions: F1, F2, F3)
# ============================================================
def build_form_i(data, output_dir, folder_num=1):
    folder_names = {
        1: ("Folder 1: Pre-requisites for Bid Acceptance", "الملف الأول – المتطلبات المسبقة لقبول العطاء"),
        2: ("Folder 2: Technical Submission", "الملف الثاني – العرض الفني"),
        3: ("Folder 3: Commercial Submission", "الملف الثالث – العرض المالي"),
    }
    
    folder_items = {
        1: [
            ("Proposal Checklist", "قائمة تدقيق العطاء", "Vol. IV – Form I"),
            ("Proposal Cover Letter", "خطاب العطاء المقدم", "Vol. IV – Form A"),
            ("Bidder's Experience", "خبرة مقدم العطاء", "Vol. IV – Form D"),
            ("Bidder's Details (and required attachments)", "بيانات مقدم العطاء والمرفقات", "Vol. IV – Form D"),
            ("Sworn Declaration", "إعلان محلف", "Vol. IV – Form E"),
            ("Trade License", "رخصة تجارية", "Vol. II – Section 5"),
            ("Manager's Cheque", "شيك مدير", "Vol. II – Section 5"),
        ],
        2: [
            ("Proposal Checklist", "قائمة تدقيق العطاء", "Vol. IV – Form I"),
            ("Bidder's Experience in Detail", "تفاصيل الخبرة السابقة", "Vol. IV – Form G"),
            ("Architectural Firm / Designer Experience", "تفاصيل خبرة شركة الهندسة", "Vol. IV – Form H"),
            ("Design Proposals", "التصميم المقترح", "Vol. VI – App. E"),
            ("Development Cost", "تكاليف التطوير", "Vol. VI – App. E"),
            ("ESG Policy", "سياسة النهج البيئي والاجتماعي والحوكمة", "Vol. VI – App. E"),
        ],
        3: [
            ("Proposal Checklist", "قائمة تدقيق العطاء", "Vol. IV – Form I"),
            ("Bidder's Proposed Musataha Fee", "قيمة المساطحة المقترحة", "Vol. IV – Form B / Vol. VI – App. F"),
            ("Feasibility Study", "دراسة الجدوى المالية", "Vol. VI – App. F / G"),
            ("Project Financing Capability", "القدرة على تمويل المشروع", "Vol. VI – App. F / G"),
        ],
    }
    
    doc = Document()
    setup_page(doc)
    add_page_footer(doc, data.get('company_legal_name', ''))
    
    title_en, title_ar = folder_names[folder_num]
    add_header_bar(doc, "FORM I", f"Proposal Checklist - {title_en}", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": f"Proposal Checklist — Form I (Folder {folder_num})"}
    add_meta_block(doc, meta)
    
    add_section_header(doc, title_en.upper())
    
    items = folder_items[folder_num]
    table = doc.add_table(rows=len(items) + 1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ["#", "Description", "الوصف", "Reference", "Submitted?"]
    widths = [Cm(1), Cm(5), Cm(4.5), Cm(3.5), Cm(2)]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.width = widths[i]
        set_cell_background(c, "0B3D7A")
        set_cell_borders(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    
    for i, (desc_en, desc_ar, ref) in enumerate(items, start=1):
        cells_data = [str(i), desc_en, desc_ar, ref, "☐ Yes / ☐ No"]
        for j, val in enumerate(cells_data):
            c = table.rows[i].cells[j]
            c.width = widths[j]
            set_cell_borders(c)
            if i % 2 == 0:
                set_cell_background(c, "F5F5F5")
            
            p = c.paragraphs[0]
            if j == 2:  # Arabic
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif j == 4:  # Submitted
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = p.add_run(str(val))
            run.font.size = Pt(9)
    
    doc.add_paragraph()
    add_signature_block(doc, data)
    
    out = Path(output_dir) / f"Form_I_Checklist_Folder{folder_num}.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# FORM G — Bidder's Experience in Detail
# ============================================================
def build_form_g(data, output_dir):
    doc = Document()
    
    # Landscape for wide table
    setup_page(doc, landscape=True)
    add_page_footer(doc, data.get('company_legal_name', ''))
    
    add_header_bar(doc, "FORM G", "Bidder's Experience in Detail", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": "Form G — Bidder's Experience in Detail"}
    add_meta_block(doc, meta)
    
    add_section_header(doc, "COMPLETED PROJECTS PORTFOLIO")
    
    projects = data.get("projects_completed", [])
    headers = ["#", "Project", "Scope", "Started", "Completed", "Amount (AED)", "Status", "Dev", "Leasing", "M&O", "GFA (m²)", "GLA (m²)", "Floor Eff.", "Occ. 2022", "Occ. 2023"]
    
    table = doc.add_table(rows=len(projects) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_background(c, "0B3D7A")
        set_cell_borders(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)
    
    # Data
    if not projects:
        # Empty template — add 8 empty rows
        for i in range(1, 9):
            row = table.add_row() if i > len(projects) else None
        # Actually we created the table with len(projects)+1 rows already, so add empties:
        for _ in range(8):
            table.add_row()
        for i in range(1, 9):
            for j in range(len(headers)):
                c = table.rows[i].cells[j]
                set_cell_borders(c)
                if i % 2 == 0:
                    set_cell_background(c, "F5F5F5")
                c.paragraphs[0].add_run(str(i) if j == 0 else "_____").font.size = Pt(8)
    else:
        for i, p in enumerate(projects, start=1):
            vals = [
                str(i),
                p.get("name", ""),
                p.get("scope", ""),
                p.get("start", ""),
                p.get("end", ""),
                f"{p.get('amount', 0):,}" if isinstance(p.get('amount'), (int, float)) else str(p.get('amount', '')),
                p.get("status", "Built 100%"),
                p.get("dev_role", "YES"),
                p.get("leasing_role", "YES"),
                p.get("mgmt_role", "YES"),
                str(p.get("gfa", "")),
                str(p.get("gla", "")),
                str(p.get("floor_eff", "")),
                str(p.get("occ_2022", "")),
                str(p.get("occ_2023", "")),
            ]
            for j, v in enumerate(vals):
                c = table.rows[i].cells[j]
                set_cell_borders(c)
                if i % 2 == 0:
                    set_cell_background(c, "F5F5F5")
                run = c.paragraphs[0].add_run(str(v))
                run.font.size = Pt(8)
    
    doc.add_paragraph()
    
    # Summary
    add_section_header(doc, "PORTFOLIO SUMMARY")
    summary_table = doc.add_table(rows=5, cols=2)
    summary = [
        ("Total Completed Projects", str(len(projects)) if projects else "_____"),
        ("Total Investment (AED)", f"AED {sum(p.get('amount', 0) for p in projects if isinstance(p.get('amount'), (int, float))):,}" if projects else "_____"),
        ("Total GFA (m²)", str(sum(p.get('gfa', 0) for p in projects if isinstance(p.get('gfa'), (int, float)))) if projects else "_____"),
        ("Total GLA (m²)", str(sum(p.get('gla', 0) for p in projects if isinstance(p.get('gla'), (int, float)))) if projects else "_____"),
        ("Years of Experience", data.get("years_experience", "_____")),
    ]
    for i, (k, v) in enumerate(summary):
        c1 = summary_table.rows[i].cells[0]
        c2 = summary_table.rows[i].cells[1]
        set_cell_background(c1, "0B3D7A")
        set_cell_borders(c1); set_cell_borders(c2)
        run1 = c1.paragraphs[0].add_run(k)
        run1.bold = True; run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run1.font.size = Pt(10)
        run2 = c2.paragraphs[0].add_run(str(v))
        run2.font.size = Pt(10)
    
    add_signature_block(doc, data)
    
    out = Path(output_dir) / "Form_G_Experience_Detail.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# FORM H — Architectural Firm Experience
# ============================================================
def build_form_h(data, output_dir):
    doc = Document()
    setup_page(doc)
    add_page_footer(doc, data.get('company_legal_name', ''))
    
    add_header_bar(doc, "FORM H", "Architectural Firm / Designer Experience", logo_path=data.get("logo_path"))
    add_recipient_block(doc)
    
    meta = {**data, "subject": "Form H — Architectural Firm Experience"}
    add_meta_block(doc, meta)
    
    add_section_header(doc, "ARCHITECTURAL FIRM PROFILE")
    
    arch = data.get("architect", {})
    profile_data = [
        ("Firm Name", arch.get("firm_name", "________________")),
        ("Contact Person", arch.get("contact", "________________")),
        ("Phone", arch.get("phone", "________________")),
        ("Email", arch.get("email", "________________")),
        ("License No. (Abu Dhabi)", arch.get("license", "________________")),
        ("Years in Operation", arch.get("years", "________________")),
        ("Specialization", arch.get("specialty", "________________")),
    ]
    profile_table = doc.add_table(rows=len(profile_data), cols=2)
    for i, (k, v) in enumerate(profile_data):
        c1 = profile_table.rows[i].cells[0]
        c2 = profile_table.rows[i].cells[1]
        c1.width = Cm(5); c2.width = Cm(11)
        set_cell_background(c1, "0B3D7A")
        set_cell_borders(c1); set_cell_borders(c2)
        run1 = c1.paragraphs[0].add_run(k); run1.bold = True
        run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); run1.font.size = Pt(10)
        run2 = c2.paragraphs[0].add_run(str(v)); run2.font.size = Pt(10)
    
    doc.add_paragraph()
    add_section_header(doc, "DESIGNER'S PROJECT PORTFOLIO")
    
    arch_projects = data.get("architect_projects", [])
    headers = ["#", "Project Name", "Description", "GFA (m²)", "Cost (AED)", "Location", "Year", "Role"]
    
    n_rows = max(len(arch_projects), 5) + 1
    table = doc.add_table(rows=n_rows, cols=len(headers))
    
    # Header
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_background(c, "0B3D7A")
        set_cell_borders(c)
        run = c.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    
    # Data
    for i in range(1, n_rows):
        proj = arch_projects[i-1] if i-1 < len(arch_projects) else {}
        vals = [
            str(i),
            proj.get("name", "_____"),
            proj.get("description", "_____"),
            str(proj.get("gfa", "_____")),
            f"AED {proj.get('cost', 0):,}" if isinstance(proj.get('cost'), (int, float)) else str(proj.get("cost", "_____")),
            proj.get("location", "_____"),
            str(proj.get("year", "_____")),
            proj.get("role", "Designer"),
        ]
        for j, v in enumerate(vals):
            c = table.rows[i].cells[j]
            set_cell_borders(c)
            if i % 2 == 0:
                set_cell_background(c, "F5F5F5")
            c.paragraphs[0].add_run(str(v)).font.size = Pt(9)
    
    add_signature_block(doc, data)
    
    out = Path(output_dir) / "Form_H_Architect_Experience.docx"
    doc.save(str(out))
    print(f"  ✓ {out.name}")
    return out


# ============================================================
# MAIN
# ============================================================
def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)
    
    data_path = Path(sys.argv[1])
    out_dir = Path(sys.argv[2])
    out_dir.mkdir(parents=True, exist_ok=True)
    
    data = json.loads(data_path.read_text(encoding='utf-8'))
    
    print(f"\n🏢 Building ADIO Tender Forms for: {data.get('company_legal_name', '?')}")
    print(f"📋 Tender: {data.get('tender_name', '?')}\n")
    
    # Build all forms
    print("📦 Package 1 — Pre-requisites:")
    build_form_i(data, out_dir, folder_num=1)
    build_form_a(data, out_dir)
    build_form_d(data, out_dir)
    build_form_e(data, out_dir)
    
    print("\n📦 Package 2 — Technical Submission:")
    build_form_i({**data, "package": 2}, out_dir, folder_num=2)
    build_form_g(data, out_dir)
    build_form_h(data, out_dir)
    
    print("\n📦 Package 3 — Commercial Submission:")
    build_form_i({**data, "package": 3}, out_dir, folder_num=3)
    
    print(f"\n✅ All forms generated in: {out_dir}")


if __name__ == "__main__":
    main()
