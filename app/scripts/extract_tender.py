#!/usr/bin/env python3
"""
extract_tender.py — استخراج تلقائي لمحتوى كل مناقصة من مجلد فُكّ من ZIP

الاستخدام:
  python3 extract_tender.py <input_dir> <output_dir>

المدخلات: مجلد فيه subfolders لكل مناقصة (زي اللي بيطلع من ZIP)
المخرجات: مجلد فيه لكل مناقصة:
  <tender_name>/
    ├── raw_text/
    │   ├── auction_document.txt
    │   ├── checklist.txt
    │   ├── long_term_lease.txt
    │   ├── form_a.txt
    │   ├── form_b.txt
    │   ├── form_h.txt
    │   ├── kyc.txt
    │   └── ndu.txt
    └── files_inventory.json
"""

import os
import sys
import json
import subprocess
import shutil
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def fix_unicode_path(p):
    """يحوّل #U0627 إلى الحرف العربي الصحيح في أسماء المجلدات."""
    import re
    return re.sub(r'#U([0-9A-Fa-f]{4})', lambda m: chr(int(m.group(1), 16)), p)


def extract_pdf_text(pdf_path, out_path):
    """pdftotext"""
    try:
        subprocess.run(
            ['pdftotext', '-layout', str(pdf_path), str(out_path)],
            check=True, capture_output=True, timeout=60
        )
        return True
    except Exception as e:
        print(f"  PDF extract failed for {pdf_path}: {e}", file=sys.stderr)
        return False


def extract_docx_text(docx_path, out_path):
    """python-docx — paragraphs + tables"""
    try:
        doc = Document(str(docx_path))
        lines = []
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for i, t in enumerate(doc.tables):
            lines.append(f"\n--- TABLE {i} ---")
            for row in t.rows:
                lines.append(" | ".join(c.text.strip() for c in row.cells))
        out_path.write_text("\n".join(lines), encoding='utf-8')
        return True
    except Exception as e:
        print(f"  DOCX extract failed for {docx_path}: {e}", file=sys.stderr)
        return False


def try_extract_rar(rar_path, out_dir):
    """جرّب 7z بالأول. RAR5 ممكن مايتفكّش."""
    try:
        result = subprocess.run(
            ['7z', 'x', '-y', f'-o{out_dir}', str(rar_path)],
            capture_output=True, text=True, timeout=120
        )
        return result.returncode == 0 or 'Everything is Ok' in result.stdout
    except Exception:
        return False


def classify_file(filename):
    """يحدد نوع الملف من اسمه."""
    lower = filename.lower()
    if 'auction document' in lower or 'auction_document' in lower:
        return 'auction_document'
    if 'submission checklist' in lower or 'submission_checklist' in lower:
        return 'checklist'
    if 'long term lease' in lower or 'lease - contract' in lower:
        return 'long_term_lease'
    if 'form a' in lower or 'letter of auction' in lower:
        return 'form_a'
    if 'form b' in lower or 'experience' in lower:
        return 'form_b'
    if 'form h' in lower or 'non conflict' in lower or 'non-conflict' in lower:
        return 'form_h'
    if 'kyc' in lower:
        return 'kyc'
    if 'ndu' in lower:
        return 'ndu'
    return 'other'


def process_tender(tender_dir, out_dir):
    """يعالج مناقصة واحدة."""
    tender_name = tender_dir.name
    fixed_name = fix_unicode_path(tender_name)
    print(f"\n[+] Processing: {fixed_name}")

    safe_name = "".join(c if c.isalnum() or c in "_- " else "_" for c in fixed_name).strip()[:80]
    out_tender = out_dir / safe_name
    out_raw = out_tender / "raw_text"
    out_originals = out_tender / "originals"
    out_raw.mkdir(parents=True, exist_ok=True)
    out_originals.mkdir(parents=True, exist_ok=True)

    inventory = {
        "tender_name": fixed_name,
        "tender_name_raw": tender_name,
        "source_path": str(tender_dir),
        "files_found": [],
        "files_extracted": [],
        "rar_files": [],
        "missing": [],
    }

    # جمع كل الملفات
    all_files = []
    for root, _, files in os.walk(tender_dir):
        for f in files:
            all_files.append(Path(root) / f)

    # فُك RAR لو في
    rar_extracted_dirs = []
    for fp in all_files:
        if fp.suffix.lower() == '.rar':
            inventory["rar_files"].append(str(fp))
            rar_out = out_tender / "_rar_extracted" / fp.stem
            rar_out.mkdir(parents=True, exist_ok=True)
            if try_extract_rar(fp, rar_out):
                print(f"    ✓ Extracted RAR: {fp.name}")
                rar_extracted_dirs.append(rar_out)
            else:
                print(f"    ✗ RAR5 cannot extract: {fp.name} (sibling tender may have unpacked files)")

    # ضم الملفات اللي فُكّت من RAR للمسح
    for rd in rar_extracted_dirs:
        for root, _, files in os.walk(rd):
            for f in files:
                all_files.append(Path(root) / f)

    # تصنيف واستخراج
    found_categories = set()
    for fp in all_files:
        if fp.suffix.lower() not in ('.pdf', '.docx'):
            continue
        category = classify_file(fp.name)
        inventory["files_found"].append({
            "path": str(fp),
            "name": fp.name,
            "category": category,
            "size_kb": round(fp.stat().st_size / 1024, 1)
        })

        if category != 'other' and category not in found_categories:
            out_txt = out_raw / f"{category}.txt"
            ok = False
            if fp.suffix.lower() == '.pdf':
                ok = extract_pdf_text(fp, out_txt)
            elif fp.suffix.lower() == '.docx':
                ok = extract_docx_text(fp, out_txt)
            if ok:
                found_categories.add(category)
                inventory["files_extracted"].append({
                    "category": category,
                    "source": fp.name,
                    "output": str(out_txt.relative_to(out_tender))
                })
                # انسخ النسخة الأصلية
                ext = fp.suffix
                shutil.copy2(fp, out_originals / f"{category}{ext}")

    # حدّد اللي ناقص
    expected = ['auction_document', 'checklist', 'long_term_lease',
                'form_a', 'form_b', 'form_h', 'kyc', 'ndu']
    inventory["missing"] = [c for c in expected if c not in found_categories]

    # احفظ الـ inventory
    (out_tender / "files_inventory.json").write_text(
        json.dumps(inventory, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )
    print(f"    Found: {sorted(found_categories)}")
    if inventory["missing"]:
        print(f"    Missing: {inventory['missing']}")

    return inventory


def main():
    if len(sys.argv) != 3:
        print("Usage: extract_tender.py <input_dir> <output_dir>")
        sys.exit(1)

    input_dir = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])
    output_dir.mkdir(parents=True, exist_ok=True)

    # اعثر على مجلدات المناقصات (subfolders)
    tender_dirs = [d for d in input_dir.iterdir() if d.is_dir()]
    if not tender_dirs:
        print(f"ERROR: No tender folders found in {input_dir}")
        sys.exit(1)

    print(f"Found {len(tender_dirs)} tender folders.")

    all_inventories = []
    for td in sorted(tender_dirs):
        inv = process_tender(td, output_dir)
        all_inventories.append(inv)

    # ملخص رئيسي
    summary = {
        "total_tenders": len(all_inventories),
        "tenders": [
            {
                "name": i["tender_name"],
                "extracted_count": len(i["files_extracted"]),
                "missing": i["missing"],
                "has_rar5_blocker": bool(i["rar_files"]) and not i["files_extracted"]
            }
            for i in all_inventories
        ]
    }
    (output_dir / "MASTER_INVENTORY.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding='utf-8'
    )
    print(f"\n[✓] Done. Output: {output_dir}")
    print(f"[✓] Master inventory: {output_dir / 'MASTER_INVENTORY.json'}")


if __name__ == "__main__":
    main()
