#!/usr/bin/env python3
"""
fill_forms.py — يملأ نماذج DOCX الخاصة بالمناقصات (Form A, B, H, KYC, NDU)
من ملف JSON يحتوي على بيانات الشركة.

الاستخدام:
  python3 fill_forms.py <company_data.json> <forms_input_dir> <output_dir> [tender_meta.json]

company_data.json نموذج:
{
  "company_legal_name": "ABC Investments LLC",
  "company_legal_name_ar": "إيه بي سي للاستثمار ش.ذ.م.م",
  "trade_license_no": "CN-1234567",
  "trade_license_authority": "Abu Dhabi Department of Economic Development",
  "trade_license_expiry": "2027-05-15",
  "country_of_incorporation": "United Arab Emirates",
  "legal_form": "Limited Liability Company (LLC)",
  "date_of_incorporation": "2015-03-10",
  "vat_no": "100123456700003",
  "economic_license_no": "CN-1234567",
  "hq_address": "Office 1801, Maqam Tower, Sowwah Square, Al Maryah Island, Abu Dhabi, UAE",
  "tax_residency": "United Arab Emirates",
  "primary_contact": {"name":"...", "title":"...", "phone":"...", "email":"..."},
  "authorized_signatory": {"name":"...", "title":"...", "phone":"...", "email":"...", "emirates_id":"..."},
  "business_overview": "...",
  "ubo_list": [{"name":"...", "nationality":"...", "ownership_pct":51, "emirates_id":"..."}],
  "executives": [{"name":"...", "title":"...", "emirates_id":"..."}],
  "financials_3yr": {"revenue":[100,120,150], "profit":[10,15,20], "net_worth":500, "liabilities":200, "currency":"AED millions"},
  "projects_past_3yr": [
    {"name":"...", "location":"...", "client":"...", "client_contact":"...", "scope":"...", "value_aed":50000000, "start":"2022-01", "end":"2023-06"}
  ],
  "gov_experience": "Worked with ADM on... and DMT on...",
  "pep_status": false,
  "legal_proceedings": false,
  "blacklisted": false
}

tender_meta.json (اختياري — لتخصيص النموذج لمناقصة معينة):
{
  "auction_title": "Establishment of an Automotive Service Center, Al Shahamah",
  "auction_date": "2026-06-30",
  "bid_validity_days": 120
}
"""

import sys
import json
from pathlib import Path
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("ERROR: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def replace_in_paragraph(para, replacements):
    """يستبدل النصوص في الـ paragraph مع الحفاظ على التنسيق."""
    for old, new in replacements.items():
        if old in para.text:
            # دمج runs ثم استبدال — حل بسيط
            inline = para.runs
            full_text = para.text
            new_text = full_text.replace(old, str(new))
            # امسح كل الـ runs وحط واحد جديد
            for run in inline:
                run.text = ''
            if inline:
                inline[0].text = new_text
            else:
                para.add_run(new_text)


def fill_form_a(template_path, output_path, company, tender_meta):
    """Form A — Letter of Auction. ببساطة بنحط التاريخ، عنوان المناقصة، اسم المقدم."""
    doc = Document(str(template_path))
    
    replacements = {
        "[Date]": tender_meta.get("submission_date", "[يُملأ يوم التقديم]"),
        "[Bidder]": company.get("company_legal_name", "[اسم الشركة]"),
    }
    
    # في فقرة عنوان المناقصة عادة فاضي بعد "Auction Title:"
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
        if para.text.strip().startswith("Auction Title:"):
            # اضف اسم المناقصة
            if tender_meta.get("auction_title"):
                para.text = f"Auction Title: {tender_meta['auction_title']}"
    
    doc.save(str(output_path))
    print(f"  ✓ Saved: {output_path.name}")


def fill_form_b(template_path, output_path, company, tender_meta):
    """Form B — Experience & Capabilities. جدول المشاريع السابقة."""
    doc = Document(str(template_path))
    
    projects = company.get("projects_past_3yr", [])
    
    # الجدول الأول (Relevant Project Experience) — أول 3 صفوف فاضية
    if doc.tables and projects:
        table = doc.tables[0]
        # الصفوف من 1 لـ 3 (بعد header) للمشاريع
        for i, project in enumerate(projects[:3]):
            if i + 1 < len(table.rows):
                row = table.rows[i + 1]
                cells = row.cells
                if len(cells) >= 7:
                    cells[0].text = str(i + 1)
                    cells[1].text = f"{project.get('name','')}\n{project.get('location','')}"
                    cells[2].text = f"{project.get('client','')}\n{project.get('client_contact','')}"
                    cells[3].text = project.get('scope', '')
                    cells[4].text = f"AED {project.get('value_aed', 0):,}"
                    cells[5].text = project.get('start', '')
                    cells[6].text = project.get('end', '')
    
    doc.save(str(output_path))
    print(f"  ✓ Saved: {output_path.name}")


def fill_form_h(template_path, output_path, company, tender_meta):
    """Form H — Non-Conflict Declaration. اسم الشركة + التوقيع."""
    doc = Document(str(template_path))
    replacements = {
        "[Bidder]": company.get("company_legal_name", "[اسم الشركة]"),
        "[Date]": tender_meta.get("submission_date", "[يُملأ يوم التقديم]"),
        "[Company Name]": company.get("company_legal_name", "[اسم الشركة]"),
    }
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    doc.save(str(output_path))
    print(f"  ✓ Saved: {output_path.name}")


def fill_kyc(template_path, output_path, company):
    """KYC Form — حقول كثيرة. بنملأ ما نقدر عليه."""
    doc = Document(str(template_path))
    
    # الـKYC في الأصل مش placeholders، هي حقول بعد ":". 
    # هنروح لكل paragraph ونشوف لو فيه تطابق ونعدّل
    field_map = {
        "Full legal name of company:": company.get("company_legal_name", ""),
        "Company license number:": company.get("trade_license_no", ""),
        "Country of incorporation:": company.get("country_of_incorporation", ""),
        "Date of incorporation:": company.get("date_of_incorporation", ""),
        "Legal Form:": company.get("legal_form", ""),
        "Tax Residency": company.get("tax_residency", ""),
        "Principal place of business:": company.get("hq_address", ""),
        "Economic Licence No.:": company.get("economic_license_no", ""),
        "VAT registration No.:": company.get("vat_no", ""),
        "Registered Business Address of Headquarters:": company.get("hq_address", ""),
        "Overview of your business operations": company.get("business_overview", ""),
    }
    
    pc = company.get("primary_contact", {})
    sig = company.get("authorized_signatory", {})
    
    for para in doc.paragraphs:
        text = para.text.strip()
        for key, value in field_map.items():
            if text == key.strip() or text == key.strip().rstrip(":"):
                # ضع القيمة في نفس الـparagraph
                para.add_run(f" {value}")
                break
    
    doc.save(str(output_path))
    print(f"  ✓ Saved: {output_path.name}")
    print(f"    ⚠️  KYC has many fields - please review and complete manually before submission")


def fill_ndu(template_path, output_path, company):
    """NDU — اسم الشركة + التوقيع."""
    doc = Document(str(template_path))
    replacements = {
        "[Company Name]": company.get("company_legal_name", ""),
        "[Bidder]": company.get("company_legal_name", ""),
        "[Investor]": company.get("company_legal_name", ""),
    }
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    doc.save(str(output_path))
    print(f"  ✓ Saved: {output_path.name}")


def main():
    if len(sys.argv) < 4:
        print(__doc__)
        sys.exit(1)
    
    company_data_path = Path(sys.argv[1])
    forms_dir = Path(sys.argv[2])
    output_dir = Path(sys.argv[3])
    tender_meta_path = Path(sys.argv[4]) if len(sys.argv) > 4 else None
    
    output_dir.mkdir(parents=True, exist_ok=True)
    
    company = json.loads(company_data_path.read_text(encoding='utf-8'))
    tender_meta = {}
    if tender_meta_path and tender_meta_path.exists():
        tender_meta = json.loads(tender_meta_path.read_text(encoding='utf-8'))
    
    print(f"Company: {company.get('company_legal_name', '?')}")
    print(f"Tender: {tender_meta.get('auction_title', '?')}")
    print()
    
    # ابحث عن النماذج
    forms = {
        'form_a': None, 'form_b': None, 'form_h': None,
        'kyc': None, 'ndu': None
    }
    for f in forms_dir.rglob("*.docx"):
        name = f.name.lower()
        if 'form a' in name or 'letter of auction' in name:
            forms['form_a'] = f
        elif 'form b' in name or 'experience' in name:
            forms['form_b'] = f
        elif 'form h' in name or 'non conflict' in name or 'non-conflict' in name:
            forms['form_h'] = f
        elif 'kyc' in name:
            forms['kyc'] = f
        elif 'ndu' in name:
            forms['ndu'] = f
    
    if forms['form_a']:
        fill_form_a(forms['form_a'], output_dir / "Form_A_Letter_of_Auction_FILLED.docx", company, tender_meta)
    if forms['form_b']:
        fill_form_b(forms['form_b'], output_dir / "Form_B_Experience_Capabilities_FILLED.docx", company, tender_meta)
    if forms['form_h']:
        fill_form_h(forms['form_h'], output_dir / "Form_H_Non_Conflict_FILLED.docx", company, tender_meta)
    if forms['kyc']:
        fill_kyc(forms['kyc'], output_dir / "Investor_KYC_FILLED.docx", company)
    if forms['ndu']:
        fill_ndu(forms['ndu'], output_dir / "Investor_NDU_FILLED.docx", company)
    
    print(f"\n[✓] Forms filled in: {output_dir}")
    print("[!] Please review each form, complete remaining fields manually, sign & stamp before submission.")


if __name__ == "__main__":
    main()
