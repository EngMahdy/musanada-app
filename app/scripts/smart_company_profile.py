#!/usr/bin/env python3
"""
📄 Smart Company Profile Generator
====================================
يولّد بروفايل شركة احترافي كامل بصيغة DOCX:
- صفحة غلاف
- نبذة عن الشركة
- الخدمات
- المشاريع السابقة (من smart_project_generator)
- الفريق الفني
- شهادات وعضويات
- معلومات الاتصال

Usage:
    python3 smart_company_profile.py <company_data.json> <projects.json> <output.docx>
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


COLOR_NAVY = RGBColor(0x0C, 0x1A, 0x35)
COLOR_GOLD = RGBColor(0xC9, 0xA8, 0x4C)
COLOR_TEXT = RGBColor(0x1E, 0x29, 0x3B)
COLOR_GRAY = RGBColor(0x64, 0x74, 0x8B)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, size=18, color=COLOR_NAVY, gold_underline=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = True
    run.font.color.rgb = color
    
    if gold_underline:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        # Gold horizontal line via Unicode
        line = p2.add_run("─" * 40)
        line.font.color.rgb = COLOR_GOLD


def add_body(doc, text, size=11, color=COLOR_TEXT, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = COLOR_TEXT


def generate_company_profile(company_data: dict, projects_data: dict, output_path: str):
    """
    يولّد بروفايل كامل (15-20 صفحة)
    """
    doc = Document()
    
    # === COVER PAGE ===
    doc.add_paragraph().add_run("\n" * 5)
    
    # Company name centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_data.get("company_name_ar", "اسم الشركة"))
    run.font.name = 'Calibri'
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = COLOR_NAVY
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(company_data.get("company_name_en", "Company Name"))
    run2.font.size = Pt(18)
    run2.font.color.rgb = COLOR_GOLD
    run2.italic = True
    
    doc.add_paragraph().add_run("\n" * 3)
    
    # Subtitle
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("ملف الشركة التعريفي\nCompany Profile")
    run3.font.size = Pt(20)
    run3.bold = True
    run3.font.color.rgb = COLOR_NAVY
    
    doc.add_paragraph().add_run("\n" * 5)
    
    # License box
    license_table = doc.add_table(rows=4, cols=2)
    license_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_rows = [
        ("رقم الرخصة Trade License", company_data.get("license_no", "CN-XXXXXX")),
        ("تاريخ التأسيس Established", company_data.get("established", "2015")),
        ("نشاط الشركة Activity", company_data.get("activity", "استشارات هندسية ودراسات الجدوى")),
        ("الإمارة Emirate", "أبوظبي Abu Dhabi"),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        cells = license_table.rows[i].cells
        cells[0].text = label
        cells[1].text = str(value)
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(11)
        set_cell_bg(cells[0], "F0F4F8")
    
    doc.add_page_break()
    
    # === ABOUT US ===
    add_heading(doc, "نبذة عن الشركة | About Us")
    
    activity_desc = company_data.get("activity_description", "")
    if not activity_desc:
        activity_desc = (
            f"تأسست شركة {company_data.get('company_name_ar', 'الشركة')} في إمارة أبوظبي "
            f"عام {company_data.get('established', '2015')}، "
            "كواحدة من الشركات الرائدة في مجال الاستشارات الهندسية وإدارة المشاريع ودراسات الجدوى الاقتصادية. "
            "نمتلك خبرة تزيد عن 10 سنوات في السوق الإماراتي، ولدينا علاقات راسخة مع جميع الجهات الحكومية "
            "في أبوظبي والإمارات. نعمل وفق أعلى المعايير الدولية ESTIDAMA وISO، "
            "ونلتزم بالجودة والمصداقية في جميع مشاريعنا."
        )
    add_body(doc, activity_desc, size=11)
    
    # === VISION & MISSION ===
    add_heading(doc, "الرؤية والرسالة | Vision & Mission", size=16)
    
    add_body(doc, "الرؤية | Vision", bold=True, size=12)
    add_body(doc,
        "أن نكون الخيار الأول والشريك الموثوق للمشاريع الهندسية والتنموية في دولة الإمارات العربية المتحدة، "
        "من خلال تقديم حلول استشارية ذكية تواكب رؤية الإمارات 2031 ومستهدفات التنمية المستدامة."
    )
    
    doc.add_paragraph()
    add_body(doc, "الرسالة | Mission", bold=True, size=12)
    add_body(doc,
        "تقديم خدمات استشارية هندسية ودراسات جدوى متكاملة بأعلى معايير الجودة، "
        "مع الالتزام بالمواعيد والميزانيات، وبناء علاقات طويلة الأمد مع عملائنا، "
        "والمساهمة في تحقيق رؤية الإمارات التنموية."
    )
    
    # === SERVICES ===
    add_heading(doc, "خدماتنا | Our Services")
    
    services = company_data.get("services", [
        ("استخراج التراخيص الهندسية", "تجارية، صناعية، استشارية"),
        ("تصنيف الشركات", "مقاولات، استشارات هندسية، جميع الدرجات"),
        ("دراسات الجدوى الاقتصادية", "NPV, IRR, تحليل الحساسية"),
        ("إدارة المناقصات الحكومية", "DMT, ADIO, الموانئ، البلديات"),
        ("التصميم المعماري والإنشائي", "مباني سكنية، تجارية، صناعية"),
        ("الإشراف على التنفيذ", "إدارة المشاريع وضمان الجودة"),
        ("معاملات الأراضي والعقارات", "تخصيص، نقل ملكية، تثمين"),
        ("الاستشارات القانونية الهندسية", "عقود، نزاعات، تحكيم"),
        ("التحليل المالي والاستثماري", "نمذجة، تحليل مخاطر، تقييم استثمارات"),
    ])
    
    for svc, desc in services:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run1 = p.add_run(f"◆ {svc}: ")
        run1.bold = True
        run1.font.color.rgb = COLOR_NAVY
        run1.font.size = Pt(11)
        run2 = p.add_run(desc)
        run2.font.color.rgb = COLOR_TEXT
        run2.font.size = Pt(10)
    
    doc.add_page_break()
    
    # === PROJECTS ===
    add_heading(doc, "مشاريعنا السابقة | Past Projects")
    
    projects = projects_data.get("projects", [])
    
    # Summary table
    summary_table = doc.add_table(rows=2, cols=3)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = summary_table.rows[0].cells
    headers[0].text = "عدد المشاريع"
    headers[1].text = "إجمالي القيمة"
    headers[2].text = "القطاع الرئيسي"
    
    for c in headers:
        set_cell_bg(c, "0C1A35")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
    
    vals = summary_table.rows[1].cells
    vals[0].text = str(projects_data.get("total_projects", len(projects)))
    vals[1].text = f"AED {projects_data.get('total_value_aed', 0):,.0f}"
    vals[2].text = projects_data.get("project_type_ar", "متنوع")
    
    for c in vals:
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = COLOR_NAVY
    
    doc.add_paragraph()
    
    # Individual projects
    for i, proj in enumerate(projects, 1):
        # Project header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(10)
        run = p.add_run(f"المشروع {i}: {proj['title']}")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_NAVY
        
        # Project details table
        tbl = doc.add_table(rows=3, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        details = [
            ("الموقع", proj['location']),
            ("القيمة | المساحة", f"{proj['value_text']}  •  {proj['area_text']}"),
            ("السنة | المدة", f"{proj['year']}  •  {proj.get('duration_months', '?')} شهر"),
        ]
        
        for j, (k, v) in enumerate(details):
            cells = tbl.rows[j].cells
            cells[0].text = k
            cells[1].text = str(v)
            set_cell_bg(cells[0], "F0F4F8")
            cells[0].paragraphs[0].runs[0].bold = True
            for c in cells:
                c.paragraphs[0].runs[0].font.size = Pt(10)
        
        # Description
        if proj.get('description'):
            desc_p = doc.add_paragraph()
            desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = desc_p.add_run(proj['description'])
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = COLOR_GRAY
    
    doc.add_page_break()
    
    # === TEAM ===
    add_heading(doc, "الفريق الفني | Technical Team")
    add_body(doc, "يضم فريقنا الفني المتميز نخبة من المهندسين والاستشاريين بخبرات متنوعة:", size=11)
    
    team = [
        ("م. محمود مهدي أبوشعيشع", "المدير العام - استشاري هندسي معتمد", "10+ سنة خبرة - مهندس معماري"),
        ("م. أحمد السيد", "مدير قسم المشاريع", "8 سنوات خبرة - مهندس مدني"),
        ("م. سارة الحوسني", "استشارية تصميم معماري", "6 سنوات خبرة - مهندسة معمارية"),
        ("م. عبدالله الحوسني", "مهندس كهروميكانيكي أول", "12 سنة خبرة - MEP Engineer"),
        ("أ. خالد المنصوري", "محلل مالي واستثماري", "7 سنوات خبرة - CFA Level 2"),
    ]
    
    team_table = doc.add_table(rows=len(team) + 1, cols=3)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = team_table.rows[0].cells
    headers[0].text = "الاسم"
    headers[1].text = "الموقع الوظيفي"
    headers[2].text = "الخبرة"
    
    for c in headers:
        set_cell_bg(c, "0C1A35")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.bold = True
        run.font.size = Pt(11)
    
    for i, (name, pos, exp) in enumerate(team, 1):
        cells = team_table.rows[i].cells
        cells[0].text = name
        cells[1].text = pos
        cells[2].text = exp
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(10)
    
    # === CERTIFICATIONS ===
    add_heading(doc, "الشهادات والاعتمادات | Certifications")
    
    certs = [
        "✓ ISO 9001:2015 - نظام إدارة الجودة",
        "✓ ISO 14001:2015 - نظام الإدارة البيئية",
        "✓ ISO 45001:2018 - نظام إدارة السلامة والصحة المهنية",
        "✓ ESTIDAMA Pearl Rating - معتمد من بلدية أبوظبي",
        "✓ معتمد لدى دائرة البلديات والنقل DMT",
        "✓ معتمد لدى مكتب أبوظبي للاستثمار ADIO",
        "✓ معتمد لدى موانئ أبوظبي",
        "✓ معتمد لدى وزارة الموارد البشرية والتوطين",
        "✓ عضو جمعية المهندسين الإماراتية",
        "✓ عضو غرفة تجارة وصناعة أبوظبي ADCCI",
    ]
    
    for cert in certs:
        add_bullet(doc, cert, size=10)
    
    # === CONTACT ===
    doc.add_page_break()
    add_heading(doc, "معلومات الاتصال | Contact Information")
    
    contact_table = doc.add_table(rows=6, cols=2)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    contact_info = [
        ("📞 الهاتف", company_data.get("phone", "+971 56 966 4664")),
        ("📧 البريد الإلكتروني", company_data.get("email", "info@musanda.ae")),
        ("🌐 الموقع الإلكتروني", company_data.get("website", "www.musanda.ae")),
        ("📍 العنوان", company_data.get("address", "خليفة سيتي - أبوظبي - الإمارات العربية المتحدة")),
        ("🏢 الرخصة", company_data.get("license_no", "CN-6295947")),
        ("📅 سارية حتى", company_data.get("license_expiry", "2027/02/17")),
    ]
    
    for i, (k, v) in enumerate(contact_info):
        cells = contact_table.rows[i].cells
        cells[0].text = k
        cells[1].text = str(v)
        set_cell_bg(cells[0], "0C1A35")
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cells[0].paragraphs[0].runs[0].bold = True
        for c in cells:
            c.paragraphs[0].runs[0].font.size = Pt(11)
    
    # Save
    doc.save(output_path)
    print(f"✅ Company profile saved: {output_path}")
    return output_path


def main():
    if len(sys.argv) < 4:
        print("Usage: python3 smart_company_profile.py <company_data.json> <projects.json> <output.docx>")
        sys.exit(1)
    
    company = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    projects = json.loads(Path(sys.argv[2]).read_text(encoding="utf-8"))
    output = sys.argv[3]
    
    generate_company_profile(company, projects, output)


if __name__ == "__main__":
    main()
