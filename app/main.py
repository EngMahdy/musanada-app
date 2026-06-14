"""
Musanada Engineering Consultancy — Tender Package Generator
Production-ready FastAPI app for deployment on Render/Railway.

Architecture:
- POST /api/process → starts a background job, returns job_id immediately
- GET /api/job/{job_id} → polls job status (progress + results)
- This avoids the 5-second HTTP health-check timeout that crashes long requests
"""

import os
import sys
import json
import shutil
import zipfile
import subprocess
import uuid
import re
import threading
import time
import traceback
from datetime import datetime
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, Form, HTTPException, BackgroundTasks, Request
from fastapi.responses import FileResponse, JSONResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.base import BaseHTTPMiddleware
import uvicorn

# Max upload size: 1 GB (covers 99% of tender ZIPs)
MAX_UPLOAD_BYTES = 1024 * 1024 * 1024  # 1 GB


# ============ JOB STORE (in-memory) ============
# Maps job_id → {status, progress, stage, message, files, error, started_at, finished_at}
JOBS = {}
JOBS_LOCK = threading.Lock()


def update_job(job_id: str, **kwargs):
    """Thread-safe job update."""
    with JOBS_LOCK:
        if job_id not in JOBS:
            JOBS[job_id] = {}
        JOBS[job_id].update(kwargs)


def get_job(job_id: str) -> dict:
    """Thread-safe job read."""
    with JOBS_LOCK:
        return dict(JOBS.get(job_id, {}))

# ============ CONFIG ============
APP_DIR = Path(__file__).resolve().parent
TEMPLATES_DIR = APP_DIR / "templates"
BRAND_DIR = APP_DIR / "brand"
SCRIPTS_DIR = APP_DIR / "scripts"
DATA_DIR = APP_DIR / "data"

# On Render, use /tmp for outputs (ephemeral but writable)
OUTPUTS_DIR = Path(os.environ.get("OUTPUTS_DIR", "/tmp/musanada_outputs"))
OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)

# Apt-installed binaries
PDFTOTEXT = "pdftotext"
SEVENZ = "7z"

# ============ APP ============
app = FastAPI(title="Musanada Engineering Consultancy")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])


# Middleware to enforce upload size limit explicitly
class LimitUploadSize(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        if request.method == "POST" and request.url.path.startswith("/api/"):
            content_length = request.headers.get("content-length")
            if content_length and int(content_length) > MAX_UPLOAD_BYTES:
                return JSONResponse(
                    status_code=413,
                    content={
                        "status": "error",
                        "message": f"الملف كبير جداً. الحد الأقصى {MAX_UPLOAD_BYTES // (1024*1024)} ميجا.",
                        "size_mb": int(content_length) // (1024*1024),
                        "max_mb": MAX_UPLOAD_BYTES // (1024*1024),
                    }
                )
        return await call_next(request)

app.add_middleware(LimitUploadSize)

# Static files
app.mount("/static", StaticFiles(directory=str(BRAND_DIR)), name="static")
app.mount("/outputs", StaticFiles(directory=str(OUTPUTS_DIR)), name="outputs")


# ============ ROUTES ============
@app.get("/", response_class=HTMLResponse)
async def index():
    html = (TEMPLATES_DIR / "index.html").read_text(encoding="utf-8")
    return html


@app.get("/health")
async def health():
    return {
        "status": "ok",
        "service": "musanada-app",
        "version": "1.0.0",
        "ts": datetime.now().isoformat()
    }


@app.get("/api/job/{job_id}")
async def get_job_status(job_id: str):
    """Poll job status. Returns: status, stage, progress %, files (if done), error."""
    job = get_job(job_id)
    if not job:
        return JSONResponse({"status": "not_found", "message": "Job not found"}, status_code=404)
    return JSONResponse(job)


@app.post("/api/process")
async def process_tender(
    background_tasks: BackgroundTasks,
    # Tender files (multiple allowed)
    tender_file: list[UploadFile] = File(...),
    # Company data — ALL OPTIONAL (AI will extract from uploaded documents if not provided)
    company_legal_name: str = Form(""),
    company_short_name: str = Form(""),
    legal_form: str = Form(""),
    establishment_date: str = Form(""),
    nature_of_business: str = Form(""),
    hq_address: str = Form(""),
    company_phone: str = Form(""),
    company_fax: str = Form(""),
    company_email: str = Form(""),
    company_website: str = Form(""),
    trade_license_no: str = Form(""),
    authorized_signatory_name: str = Form(""),
    authorized_signatory_title: str = Form(""),
    contact_person_name: str = Form(""),
    contact_person_phone: str = Form(""),
    contact_person_email: str = Form(""),
    years_experience: str = Form(""),
    experience_domain: str = Form(""),
    logo: UploadFile = File(None),
    signature: UploadFile = File(None),
    stamp: UploadFile = File(None),
    # Company supporting documents — ALL accept multiple files
    doc_license: list[UploadFile] = File([]),
    doc_profile: list[UploadFile] = File([]),
    doc_works: list[UploadFile] = File([]),
    doc_financials: list[UploadFile] = File([]),
    doc_bank: list[UploadFile] = File([]),
    doc_id: list[UploadFile] = File([]),
    doc_orgchart: list[UploadFile] = File([]),
    doc_other: list[UploadFile] = File([]),
    proj_name: list = Form([]),
    proj_location: list = Form([]),
    proj_scope: list = Form([]),
    proj_amount: list = Form([]),
    proj_start: list = Form([]),
    proj_end: list = Form([]),
    proj_gfa: list = Form([]),
):
    """Submit tender for processing. Returns job_id immediately - poll /api/job/{id}.
    
    IMPORTANT: This handler must finish QUICKLY (< 30s) to avoid HTTP timeout.
    All file saves happen synchronously, then heavy work is spawned in a thread.
    """
    job_id = uuid.uuid4().hex[:12]
    job_dir = OUTPUTS_DIR / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    
    # Initialize job state
    update_job(job_id,
        status="processing",
        stage="رفع الملفات...",
        progress=5,
        started_at=datetime.now().isoformat(),
    )
    
    try:
        # ===== STEP 1: Save tender file(s) =====
        tender_input_dir = job_dir / "tender_input"
        tender_input_dir.mkdir(exist_ok=True)
        
        tender_files_saved = []
        for uf in tender_file:
            if uf and uf.filename:
                safe_name = uf.filename.replace("/", "_").replace("\\", "_")
                p = tender_input_dir / safe_name
                with open(p, "wb") as f:
                    shutil.copyfileobj(uf.file, f)
                tender_files_saved.append(p)
                print(f"✓ Saved tender file: {safe_name} ({p.stat().st_size//1024} KB)")
        
        if not tender_files_saved:
            raise HTTPException(400, "لم يتم رفع أي ملف للمناقصة")
        
        # Primary file (first uploaded — used for the legacy single-path logic)
        tender_path = tender_files_saved[0]
        tender_filename = tender_path.name
        
        # ===== STEP 2: Save images =====
        images_dir = job_dir / "images"
        images_dir.mkdir(exist_ok=True)
        logo_path = None
        sig_path = None
        stamp_path = None
        
        if logo and logo.filename:
            logo_path = images_dir / f"logo{Path(logo.filename).suffix}"
            with open(logo_path, "wb") as f:
                shutil.copyfileobj(logo.file, f)
        
        if signature and signature.filename:
            sig_path = images_dir / f"signature{Path(signature.filename).suffix}"
            with open(sig_path, "wb") as f:
                shutil.copyfileobj(signature.file, f)
        
        if stamp and stamp.filename:
            stamp_path = images_dir / f"stamp{Path(stamp.filename).suffix}"
            with open(stamp_path, "wb") as f:
                shutil.copyfileobj(stamp.file, f)
        
        # ===== STEP 2.5: Save company documents (all support multi-file) =====
        docs_dir = job_dir / "company_docs"
        docs_dir.mkdir(exist_ok=True)
        
        saved_docs = {}  # {category: [paths]}
        all_doc_categories = {
            "license": doc_license,
            "profile": doc_profile,
            "works": doc_works,
            "financials": doc_financials,
            "bank": doc_bank,
            "id": doc_id,
            "orgchart": doc_orgchart,
            "other": doc_other,
        }
        for category, file_list in all_doc_categories.items():
            if not file_list:
                continue
            cat_dir = docs_dir / category
            cat_dir.mkdir(exist_ok=True)
            for uf in file_list:
                if uf and uf.filename:
                    safe_name = uf.filename.replace("/", "_").replace("\\", "_")
                    p = cat_dir / safe_name
                    # Handle duplicate names
                    counter = 1
                    while p.exists():
                        stem = Path(safe_name).stem
                        ext = Path(safe_name).suffix
                        p = cat_dir / f"{stem}_{counter}{ext}"
                        counter += 1
                    with open(p, "wb") as f:
                        shutil.copyfileobj(uf.file, f)
                    saved_docs.setdefault(category, []).append(p)
                    print(f"✓ Saved {category}: {p.name}")
        
        # ===== Upload phase done — spawn background thread =====
        update_job(job_id, stage="جاري المعالجة في الخلفية...", progress=15)
        
        # Snapshot user-provided form data into a dict for the background worker
        form_snapshot = {
            "company_legal_name": company_legal_name,
            "company_short_name": company_short_name,
            "legal_form": legal_form,
            "establishment_date": establishment_date,
            "nature_of_business": nature_of_business,
            "hq_address": hq_address,
            "company_phone": company_phone,
            "company_fax": company_fax,
            "company_email": company_email,
            "company_website": company_website,
            "trade_license_no": trade_license_no,
            "authorized_signatory_name": authorized_signatory_name,
            "authorized_signatory_title": authorized_signatory_title,
            "contact_person_name": contact_person_name,
            "contact_person_phone": contact_person_phone,
            "contact_person_email": contact_person_email,
            "years_experience": years_experience,
            "experience_domain": experience_domain,
            "proj_name": list(proj_name),
            "proj_location": list(proj_location),
            "proj_scope": list(proj_scope),
            "proj_amount": list(proj_amount),
            "proj_start": list(proj_start),
            "proj_end": list(proj_end),
            "proj_gfa": list(proj_gfa),
            "logo_path": str(logo_path) if logo_path else None,
            "sig_path": str(sig_path) if sig_path else None,
            "stamp_path": str(stamp_path) if stamp_path else None,
            "tender_files_saved": [str(p) for p in tender_files_saved],
            "saved_docs": {k: [str(p) for p in v] for k, v in saved_docs.items()},
        }
        
        # Save snapshot so background thread can read it
        (job_dir / "form_snapshot.json").write_text(
            json.dumps(form_snapshot, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        
        # Spawn background thread
        thread = threading.Thread(
            target=_run_processing_in_background,
            args=(job_id, str(job_dir)),
            daemon=True,
        )
        thread.start()
        
        # Return immediately with job_id
        return JSONResponse({
            "status": "processing",
            "job_id": job_id,
            "poll_url": f"/api/job/{job_id}",
            "message": "تم استلام الملفات. تابع التقدم عبر poll_url"
        })
    
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        update_job(job_id, status="error", message=str(e), error=traceback.format_exc()[:1000])
        return JSONResponse({"status": "error", "job_id": job_id, "message": str(e)}, status_code=500)


def _run_processing_in_background(job_id: str, job_dir_str: str):
    """The heavy-lifting function. Runs in a background thread.
    Updates JOBS[job_id] with progress."""
    job_dir = Path(job_dir_str)
    
    try:
        # Load form snapshot
        form_data = json.loads((job_dir / "form_snapshot.json").read_text(encoding="utf-8"))
        
        # Reconstruct paths
        saved_docs = {k: [Path(p) for p in v] for k, v in form_data["saved_docs"].items()}
        tender_files_saved = [Path(p) for p in form_data["tender_files_saved"]]
        logo_path = Path(form_data["logo_path"]) if form_data.get("logo_path") else None
        sig_path = Path(form_data["sig_path"]) if form_data.get("sig_path") else None
        stamp_path = Path(form_data["stamp_path"]) if form_data.get("stamp_path") else None
        
        # Reconstruct user input vars (just aliases)
        company_legal_name = form_data["company_legal_name"]
        company_short_name = form_data["company_short_name"]
        legal_form = form_data["legal_form"]
        establishment_date = form_data["establishment_date"]
        nature_of_business = form_data["nature_of_business"]
        hq_address = form_data["hq_address"]
        company_phone = form_data["company_phone"]
        company_fax = form_data["company_fax"]
        company_email = form_data["company_email"]
        company_website = form_data["company_website"]
        trade_license_no = form_data["trade_license_no"]
        authorized_signatory_name = form_data["authorized_signatory_name"]
        authorized_signatory_title = form_data["authorized_signatory_title"]
        contact_person_name = form_data["contact_person_name"]
        contact_person_phone = form_data["contact_person_phone"]
        contact_person_email = form_data["contact_person_email"]
        years_experience = form_data["years_experience"]
        experience_domain = form_data["experience_domain"]
        proj_name = form_data["proj_name"]
        proj_location = form_data["proj_location"]
        proj_scope = form_data["proj_scope"]
        proj_amount = form_data["proj_amount"]
        proj_start = form_data["proj_start"]
        proj_end = form_data["proj_end"]
        proj_gfa = form_data["proj_gfa"]
        
        # Primary tender filename
        tender_path = tender_files_saved[0]
        tender_filename = tender_path.name
        
        # ===== AI extraction =====
        update_job(job_id, stage="استخراج البيانات من الملفات (AI)...", progress=20)
        
        # Extract intelligence from documents (process ALL files in each category)
        extracted_intel = {}
        if saved_docs.get("license"):
            # Use first license file (usually only one)
            extracted_intel["license"] = extract_license_data(saved_docs["license"][0])
        if saved_docs.get("profile"):
            # Combine all profile files
            extracted_intel["profile"] = extract_profile_data(saved_docs["profile"][0])
        if saved_docs.get("works"):
            # Extract projects from ALL works files & merge
            all_projects = []
            for wf in saved_docs["works"]:
                try:
                    projs = extract_projects_data(wf)
                    all_projects.extend(projs)
                except Exception as e:
                    print(f"Failed to extract from {wf.name}: {e}")
            # Dedupe by name
            seen = set()
            unique_projects = []
            for p in all_projects:
                key = p.get("name", "").lower().strip()
                if key and key not in seen:
                    seen.add(key)
                    unique_projects.append(p)
            extracted_intel["projects"] = unique_projects
        if saved_docs.get("financials"):
            extracted_intel["financials"] = extract_financials_data(saved_docs["financials"][0])
        if saved_docs.get("bank"):
            extracted_intel["bank"] = extract_bank_data(saved_docs["bank"][0])
        
        # Save intel for transparency
        (job_dir / "extracted_intel.json").write_text(
            json.dumps(extracted_intel, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
        
        # ===== STEP 3: Build bidder_data.json =====
        projects = []
        for i in range(len(proj_name)):
            name = proj_name[i] if i < len(proj_name) else ""
            if not name.strip():
                continue
            projects.append({
                "name": name,
                "location": proj_location[i] if i < len(proj_location) else "",
                "scope": proj_scope[i] if i < len(proj_scope) else "",
                "amount": int(proj_amount[i]) if i < len(proj_amount) and proj_amount[i] else 0,
                "start": proj_start[i] if i < len(proj_start) else "",
                "end": proj_end[i] if i < len(proj_end) else "",
                "gfa": int(proj_gfa[i]) if i < len(proj_gfa) and proj_gfa[i] else 0,
                "gla": int(int(proj_gfa[i]) * 0.92) if i < len(proj_gfa) and proj_gfa[i] else 0,
                "status": "Built 100%",
                "dev_role": "YES", "leasing_role": "YES", "mgmt_role": "YES",
                "floor_eff": "92%", "occ_2022": "85%", "occ_2023": "90%",
            })
        
        # Merge AI-extracted data: form input takes priority, AI fills gaps
        lic_intel = extracted_intel.get("license", {})
        prof_intel = extracted_intel.get("profile", {})
        ai_projects = extracted_intel.get("projects", [])
        
        # Add AI-extracted projects to user-provided list (de-dup by name)
        existing_names = {p["name"].lower() for p in projects if p.get("name")}
        for ap in ai_projects:
            if ap.get("name", "").lower() not in existing_names:
                projects.append(ap)
                existing_names.add(ap.get("name", "").lower())
        
        # AI-first merge: form input > AI extraction > placeholder
        # User-provided values take priority, but if blank, fall back to AI-extracted data
        def best(user_val, ai_val, placeholder=""):
            """Return user_val if non-empty, else ai_val, else placeholder."""
            if user_val and str(user_val).strip():
                return user_val
            if ai_val and str(ai_val).strip():
                return ai_val
            return placeholder
        
        final_company_name = best(company_legal_name, lic_intel.get("trade_name_en"), "[Company Name]")
        final_legal_form = best(legal_form, lic_intel.get("legal_form"), "Limited Liability Company")
        final_establishment = best(establishment_date, lic_intel.get("establishment_date"), "")
        final_nature = best(nature_of_business, lic_intel.get("activities"), "")
        final_license_no = best(trade_license_no, lic_intel.get("license_no"), "")
        final_hq = best(hq_address, prof_intel.get("address") or lic_intel.get("address"), "Abu Dhabi, UAE")
        final_phone = best(company_phone, prof_intel.get("phone"), "")
        final_email = best(company_email, prof_intel.get("email"), "")
        final_website = best(company_website, prof_intel.get("website"), "")
        final_signatory = best(authorized_signatory_name, lic_intel.get("owner_name"), "[Authorized Signatory]")
        final_title = best(authorized_signatory_title, "", "Owner / Managing Director")
        
        bidder_data = {
            "company_legal_name": final_company_name,
            "company_short_name": best(company_short_name, prof_intel.get("short_name")),
            "legal_form": final_legal_form,
            "establishment_date": final_establishment,
            "nature_of_business": final_nature,
            "hq_address": final_hq,
            "bidder_address": final_hq,
            "partners_nationality": lic_intel.get("partners_nationality", "UAE"),
            "company_phone": final_phone,
            "company_fax": company_fax,
            "company_email": final_email,
            "company_website": final_website,
            "trade_license_no": final_license_no,
            "authorized_signatory_name": final_signatory,
            "authorized_signatory_title": final_title,
            "contact_person_name": best(contact_person_name, final_signatory),
            "contact_person_phone": best(contact_person_phone, final_phone),
            "contact_person_email": best(contact_person_email, final_email),
            "min_exp_years": 2,
            "experience_domain": best(experience_domain, "", "building and operating similar facilities"),
            "years_experience": years_experience,
            "submission_date": datetime.now().strftime("%d-%m-%Y"),
            "projects_completed": projects,
            "projects_current": [],
            "architect": {"firm_name": "Selected Architectural Firm", "specialty": "Commercial Architecture"},
            "architect_projects": [],
            "logo_path": str(logo_path) if logo_path else None,
            "signature_image_path": str(sig_path) if sig_path else None,
            "stamp_image_path": str(stamp_path) if stamp_path else None,
            "_extracted_intel": extracted_intel,  # for transparency
            "_data_sources": {  # show user where each field came from
                "company_legal_name": "user" if company_legal_name else ("ai_license" if lic_intel.get("trade_name_en") else "placeholder"),
                "establishment_date": "user" if establishment_date else ("ai_license" if lic_intel.get("establishment_date") else "missing"),
                "trade_license_no": "user" if trade_license_no else ("ai_license" if lic_intel.get("license_no") else "missing"),
                "authorized_signatory_name": "user" if authorized_signatory_name else ("ai_license" if lic_intel.get("owner_name") else "placeholder"),
            },
        }
        
        bidder_data_path = job_dir / "bidder_data.json"
        bidder_data_path.write_text(json.dumps(bidder_data, ensure_ascii=False, indent=2), encoding="utf-8")
        
        update_job(job_id, stage="فك ملف المناقصة (ZIP/RAR)...", progress=30)
        # ===== STEP 4: Extract & normalize tender files =====
        extracted_dir = job_dir / "extracted_tender"
        extracted_dir.mkdir(exist_ok=True)
        
        # Process ALL uploaded tender files
        all_extracted_files = []  # list of all PDFs/DOCX after extraction
        
        for tf in tender_files_saved:
            suffix = tf.suffix.lower()
            
            if suffix == '.zip':
                # Extract ZIP
                try:
                    with zipfile.ZipFile(tf, 'r') as z:
                        z.extractall(extracted_dir)
                        print(f"✓ Extracted ZIP: {tf.name}")
                except (zipfile.BadZipFile, OSError) as e:
                    print(f"⚠ Bad ZIP {tf.name}: {e}")
                    # Treat as PDF
                    shutil.copy2(tf, extracted_dir / tf.name)
            
            elif suffix == '.rar':
                # Try multiple RAR extraction methods. RAR5 needs unrar-free.
                extracted_rar = False
                rar_out_dir = extracted_dir / f"_rar_{tf.stem}"
                rar_out_dir.mkdir(exist_ok=True)
                
                # Order matters: unrar-free works on RAR5 with Unsupported Methods
                for cmd_attempt in [
                    ["unrar-free", "x", str(tf), str(rar_out_dir) + "/"],
                    ["unar", "-o", str(rar_out_dir), str(tf)],
                    ["unrar", "x", "-y", str(tf), str(rar_out_dir) + "/"],
                    ["7z", "x", "-y", f"-o{rar_out_dir}", str(tf)],
                ]:
                    try:
                        result = subprocess.run(cmd_attempt, capture_output=True, text=True, timeout=180)
                        # Verify extraction actually produced non-empty files
                        non_empty_files = [
                            p for p in rar_out_dir.rglob("*")
                            if p.is_file() and p.stat().st_size > 100
                        ]
                        if non_empty_files:
                            extracted_rar = True
                            print(f"✓ Extracted RAR with {cmd_attempt[0]}: {len(non_empty_files)} non-empty files")
                            break
                        else:
                            # Clean up empty extraction and try next tool
                            print(f"  {cmd_attempt[0]}: extracted but files are empty, trying next...")
                            for p in list(rar_out_dir.rglob("*")):
                                if p.is_file():
                                    p.unlink()
                    except (FileNotFoundError, subprocess.TimeoutExpired) as e:
                        print(f"  {cmd_attempt[0]}: {type(e).__name__}, trying next...")
                        continue
                
                if extracted_rar:
                    # Move ALL extracted contents (preserving structure) to extracted_dir
                    # Walk into rar_out_dir and move each item to extracted_dir
                    for item in list(rar_out_dir.iterdir()):
                        target = extracted_dir / item.name
                        if target.exists():
                            # Handle conflict - rename
                            counter = 1
                            while target.exists():
                                target = extracted_dir / f"{item.stem}_{counter}{item.suffix}"
                                counter += 1
                        shutil.move(str(item), str(target))
                    shutil.rmtree(rar_out_dir, ignore_errors=True)
                    print(f"  Moved RAR contents to root")
                else:
                    # RAR5 blocker: keep the file as-is (will be reported in attachments)
                    print(f"⚠ RAR5 blocker — could not extract {tf.name}, will use as attachment only")
                    rar_blocker_dir = extracted_dir / "_RAR_Could_Not_Extract"
                    rar_blocker_dir.mkdir(exist_ok=True)
                    shutil.copy2(tf, rar_blocker_dir / tf.name)
            
            elif suffix in ['.pdf', '.docx', '.doc', '.xlsx', '.xls']:
                # Copy directly (single file)
                shutil.copy2(tf, extracted_dir / tf.name)
                print(f"✓ Copied: {tf.name}")
            else:
                # Unknown - copy as-is
                shutil.copy2(tf, extracted_dir / tf.name)
                print(f"? Unknown type, copied: {tf.name}")
        
        # ===== STEP 5: Normalize structure for extract_tender.py =====
        # The script expects: input_dir/<TenderName>/Header Attachments/*.pdf
        workspace_dir = job_dir / "workspace"
        workspace_dir.mkdir(exist_ok=True)
        
        # Find all PDFs in extracted_dir (any depth)
        all_pdfs = list(extracted_dir.rglob("*.pdf")) + list(extracted_dir.rglob("*.PDF"))
        all_docs = (list(extracted_dir.rglob("*.docx")) + list(extracted_dir.rglob("*.DOCX")))
        
        print(f"=== Found {len(all_pdfs)} PDFs and {len(all_docs)} DOCX files ===")
        for p in (all_pdfs + all_docs)[:30]:
            print(f"  {p.relative_to(extracted_dir)}")
        
        # CRITICAL FIX: Check if files contain "auction document" keyword
        # If NOT, rename one to make the classifier happy
        has_auction_doc = any(
            ('auction' in p.name.lower() or 'rfp' in p.name.lower())
            for p in all_pdfs
        )
        
        if not has_auction_doc and all_pdfs:
            # Pick the largest PDF (likely the main tender document)
            largest_pdf = max(all_pdfs, key=lambda p: p.stat().st_size)
            new_name = f"Auction Document - {largest_pdf.name}"
            new_path = largest_pdf.parent / new_name
            shutil.move(str(largest_pdf), str(new_path))
            print(f"✓ Renamed largest PDF to make it the Auction Document: {new_name}")
            # Refresh list
            all_pdfs = list(extracted_dir.rglob("*.pdf")) + list(extracted_dir.rglob("*.PDF"))
        
        # Now check structure: ensure there's at least one tender subfolder with Header Attachments
        # Strategy: if we don't have a clean subfolder structure, wrap everything
        has_subdirs_with_pdfs = False
        for sub in extracted_dir.iterdir():
            if sub.is_dir() and not sub.name.startswith("_"):
                # Check if this subfolder has PDFs anywhere
                if list(sub.rglob("*.pdf")) or list(sub.rglob("*.PDF")):
                    has_subdirs_with_pdfs = True
                    # Ensure it has a "Header Attachments" subfolder
                    pdfs_at_top = list(sub.glob("*.pdf")) + list(sub.glob("*.PDF"))
                    if pdfs_at_top and not (sub / "Header Attachments").exists():
                        header_dir = sub / "Header Attachments"
                        header_dir.mkdir(exist_ok=True)
                        for pdf in pdfs_at_top:
                            shutil.move(str(pdf), str(header_dir / pdf.name))
                        print(f"✓ Wrapped PDFs in Header Attachments for {sub.name}")
        
        if not has_subdirs_with_pdfs and all_pdfs:
            # Loose files at root - wrap them
            tender_name = "Uploaded_Tender"
            if tender_files_saved:
                base = Path(tender_files_saved[0].name).stem
                tender_name = base.replace(" ", "_")[:60] or tender_name
            
            single = extracted_dir / tender_name
            single.mkdir(exist_ok=True)
            header_dir = single / "Header Attachments"
            header_dir.mkdir(exist_ok=True)
            
            # Move all loose PDFs/DOCX to Header Attachments
            for f in list(extracted_dir.iterdir()):
                if f.is_file() and f.suffix.lower() in ['.pdf', '.docx', '.doc']:
                    shutil.move(str(f), str(header_dir / f.name))
            print(f"✓ Wrapped loose files into folder: {tender_name}/Header Attachments/")
        
        input_dir = extracted_dir
        
        # Log structure for debugging
        print(f"=== Input structure ===")
        for p in sorted(extracted_dir.rglob("*")):
            if p.is_file():
                print(f"  {p.relative_to(extracted_dir)}")
        
        update_job(job_id, stage="قراءة وثيقة المناقصة...", progress=35)
        run_cmd(["python3", str(SCRIPTS_DIR / "extract_tender.py"), str(input_dir), str(workspace_dir)])
        
        update_job(job_id, stage="تحليل المناقصة بـ AI (دراسة عميقة)...", progress=45)
        # ===== STEP 6a: Run AI Reader on the main auction document PDF =====
        # Find the main Auction Document PDF in extracted_dir
        main_auction_pdf = None
        for p in extracted_dir.rglob("*.pdf"):
            name_lower = p.name.lower()
            if "auction document" in name_lower or "auction_document" in name_lower:
                main_auction_pdf = p
                break
        # Fallback: largest PDF
        if not main_auction_pdf:
            pdfs = list(extracted_dir.rglob("*.pdf"))
            if pdfs:
                main_auction_pdf = max(pdfs, key=lambda p: p.stat().st_size)
        
        tender_ai_path = job_dir / "tender_intelligence.json"
        if main_auction_pdf and os.environ.get("OPENAI_API_KEY"):
            try:
                # Use v2 (deep reader: full PDF + GPT-4o refinement + Vision)
                ai_script = SCRIPTS_DIR / "tender_ai_reader_v2.py"
                if not ai_script.exists():
                    ai_script = SCRIPTS_DIR / "tender_ai_reader.py"
                
                update_job(job_id, stage="دراسة المناقصة بالكامل بـ AI (GPT-4o + Vision)...", progress=42)
                run_cmd(
                    ["python3", "-u", str(ai_script),
                     str(main_auction_pdf), str(tender_ai_path)],
                )
                print(f"✓ AI deep intelligence saved: {tender_ai_path}")
            except Exception as ai_err:
                print(f"⚠ AI reader failed (will fall back to regex): {ai_err}")
                tender_ai_path.write_text(json.dumps({"error": str(ai_err)}), encoding="utf-8")
        else:
            print(f"⚠ No PDF or OPENAI_API_KEY — skipping AI reader")
            tender_ai_path.write_text(json.dumps({"error": "no AI"}), encoding="utf-8")
        
        update_job(job_id, stage="تحليل المناقصة (regex)...", progress=50)
        # ===== STEP 6b: Build analysis (regex - lightweight) =====
        run_cmd(["python3", str(SCRIPTS_DIR / "build_analysis.py"), str(workspace_dir)])
        
        # Find tender dirs
        tender_dirs = [d for d in workspace_dir.iterdir() 
                       if d.is_dir() and (d / "tender_meta.json").exists()]
        
        if not tender_dirs:
            # Detailed diagnostic
            workspace_contents = []
            for p in workspace_dir.rglob("*"):
                workspace_contents.append(str(p.relative_to(workspace_dir)))
            
            extracted_list = []
            for p in extracted_dir.rglob("*"):
                if p.is_file():
                    extracted_list.append(str(p.relative_to(extracted_dir)))
            
            error_detail = (
                f"تعذّر استخراج بيانات المناقصة من الملف '{tender_filename}'.\n\n"
                f"محتوى الملف المستخرج ({len(extracted_list)} ملف):\n" + 
                "\n".join(f"  - {p}" for p in extracted_list[:20]) +
                "\n\nالأسباب المحتملة:\n"
                "  1. الملف فاضي أو تالف\n"
                "  2. وثيقة المناقصة ليست PDF\n"
                "  3. صيغة Auction Document مختلفة عن DMT/ADIO المعتاد\n"
                "  4. الـPDF محمي بكلمة سر\n\n"
                "حاول رفع الملف بالكامل (ZIP فيه Auction Document.pdf)."
            )
            raise HTTPException(500, error_detail)
        
        # ===== STEP 7+: Process each tender =====
        all_outputs = []
        
        for tender_dir in tender_dirs:
            meta = json.loads((tender_dir / "tender_meta.json").read_text(encoding="utf-8"))
            tender_name = meta.get("auction_title", tender_dir.name)
            safe_name = "".join(c if c.isalnum() or c in "_- " else "_" for c in tender_name)[:60].strip()
            
            forms_data = dict(bidder_data)
            forms_data["tender_no"] = meta.get("auction_title", "")[:30]
            forms_data["tender_name"] = tender_name
            
            forms_dir = job_dir / "results" / safe_name / "01_Forms"
            forms_dir.mkdir(parents=True, exist_ok=True)
            
            forms_data_path = job_dir / f"forms_data_{safe_name}.json"
            forms_data_path.write_text(json.dumps(forms_data, ensure_ascii=False, indent=2), encoding="utf-8")
            
            # ===== Decide DMT vs ADIO based on AI intelligence =====
            try:
                intel = json.loads(tender_ai_path.read_text(encoding="utf-8"))
            except Exception:
                intel = {}
            authority = (intel.get("authority", {}).get("issuing_authority") or "").upper()
            
            # Find source DOCX files (DMT originals) in extracted dir for filling
            dmt_source_dir = None
            for d in extracted_dir.rglob("*"):
                if d.is_dir():
                    docx_files = list(d.glob("*.docx"))
                    has_form_a = any("Form A" in f.name or "Letter of Auction" in f.name for f in docx_files)
                    has_kyc = any("KYC" in f.name for f in docx_files)
                    if has_form_a or has_kyc:
                        dmt_source_dir = d
                        break
            
            # DMT path requires both AI intel AND original DOCX templates available
            use_dmt = (authority == "DMT") and dmt_source_dir and (intel.get("authority"))
            
            if use_dmt:
                update_job(job_id, stage=f"توليد النماذج DMT (Form A, B, H, KYC, NDU)...", progress=60)
                # Add intel to forms_data so DMT script can use tender title
                run_cmd(["python3", str(SCRIPTS_DIR / "fill_dmt_forms.py"),
                         str(forms_data_path),
                         str(tender_ai_path),
                         str(dmt_source_dir),
                         str(forms_dir)])
                print(f"✓ Used DMT form-filling pipeline (source: {dmt_source_dir})")
            else:
                update_job(job_id, stage=f"توليد النماذج ADIO (Form A, D, E, G, H, I)...", progress=60)
                run_cmd(["python3", str(SCRIPTS_DIR / "generate_adio_forms.py"), 
                         str(forms_data_path), str(forms_dir)])
                print(f"✓ Used ADIO form-generation pipeline (authority={authority}, dmt_dir={dmt_source_dir})")
            
            # Financial Model
            financial_dir = job_dir / "results" / safe_name / "02_Financial_Model"
            financial_dir.mkdir(parents=True, exist_ok=True)
            
            meta["tender_name"] = tender_name
            meta["issuing_authority"] = "DMT" if "DMT" in tender_name or "Auction Document" in tender_name else "ADIO"
            meta["land_area_sqm"] = float(str(meta.get("land_area_sqm", "5000")).replace(",", "") or "5000")
            meta["min_rent_per_sqm"] = float(str(meta.get("min_rent_per_sqm", "100")).replace(",", "") or "100")
            meta["contract_years"] = int(meta.get("contract_years", 25))
            meta["grace_period_years"] = int(meta.get("grace_period_years", 1))
            meta["annual_escalation_govt"] = 0.02
            (tender_dir / "tender_meta.json").write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
            
            financial_excel = financial_dir / f"Financial_Model_{safe_name}.xlsx"
            update_job(job_id, stage="بناء النموذج المالي (Excel)...", progress=70)
            run_cmd(["python3", str(SCRIPTS_DIR / "financial_model.py"),
                     str(tender_dir / "tender_meta.json"), str(financial_excel)])
            
            # Copy analysis (FIRST - cheap operation)
            analysis_dest = job_dir / "results" / safe_name / "04_Analysis"
            analysis_dest.mkdir(parents=True, exist_ok=True)
            
            analysis_src = tender_dir / "analysis_AR.md"
            if analysis_src.exists():
                shutil.copy2(analysis_src, analysis_dest / "Analysis_AR.md")
            shutil.copy2(forms_data_path, analysis_dest / "bidder_data.json")
            
            # ===== Attach company documents (BEFORE arch generation in case it fails) =====
            # Technical attachments (Package 2)
            tech_attach = job_dir / "results" / safe_name / "01_Forms" / "Attachments_Technical"
            tech_attach.mkdir(parents=True, exist_ok=True)
            
            attach_categories_tech = {
                "license": ("N8_Trade_License", "Trade License + Certificate of Incorporation"),
                "id": ("N9_Power_of_Attorney", "Power of Attorney + Emirates ID"),
                "orgchart": ("N17_Organization_Chart", "Organization Chart"),
                "works": ("N18_Bidders_Experience", "Past Projects Evidence"),
                "profile": ("Bonus_Company_Profile", "Company Profile / Brochure"),
            }
            for cat, (folder_name, desc) in attach_categories_tech.items():
                src_files = saved_docs.get(cat, [])
                if src_files:
                    target_dir = tech_attach / folder_name
                    target_dir.mkdir(exist_ok=True)
                    for src_file in src_files:
                        shutil.copy2(src_file, target_dir / src_file.name)
            
            # Financial attachments (Package 3)
            fin_attach = job_dir / "results" / safe_name / "02_Financial_Model" / "Attachments_Financial"
            fin_attach.mkdir(parents=True, exist_ok=True)
            
            attach_categories_fin = {
                "financials": ("Audited_Financial_Statements", "Audited Financials (3 years)"),
                "bank": ("Bank_Statements", "Bank Statements (Last 6 Months)"),
            }
            for cat, (folder_name, desc) in attach_categories_fin.items():
                src_files = saved_docs.get(cat, [])
                if src_files:
                    target_dir = fin_attach / folder_name
                    target_dir.mkdir(exist_ok=True)
                    for src_file in src_files:
                        shutil.copy2(src_file, target_dir / src_file.name)
            
            # Other / misc documents
            if saved_docs.get("other"):
                other_attach = job_dir / "results" / safe_name / "05_Additional_Documents"
                other_attach.mkdir(parents=True, exist_ok=True)
                for src_file in saved_docs["other"]:
                    shutil.copy2(src_file, other_attach / src_file.name)
            
            # Build attachments index
            build_attachments_index(
                job_dir / "results" / safe_name / "06_Attachments_Index.md",
                saved_docs,
                extracted_intel
            )
            
            # ===== STRATEGIC AI ADVISOR (NEW) =====
            # AI thinks strategically: pricing recommendations, ROI, market analysis,
            # per-form recommendations, action plan
            update_job(job_id, stage="🧠 AI يفكر استراتيجياً (تحليل مالي + فني + قانوني + سوقي)...", progress=72)
            try:
                advisor_md = job_dir / "results" / safe_name / "08_Strategic_Advisor.md"
                run_cmd([
                    "python3", str(SCRIPTS_DIR / "strategic_advisor.py"),
                    str(tender_ai_path),
                    str(forms_data_path),
                    str(advisor_md)
                ])
                print(f"✓ Strategic advisor saved: {advisor_md.name}")
            except Exception as adv_err:
                print(f"⚠ Strategic advisor failed: {adv_err}")
            
            # Build DYNAMIC CHECKLIST from AI intelligence
            update_job(job_id, stage="بناء قائمة المراجعة الديناميكية...", progress=75)
            try:
                # Save saved_docs as JSON for the checklist script
                saved_docs_path = job_dir / "saved_docs.json"
                saved_docs_serializable = {k: [str(p) for p in v] for k, v in saved_docs.items()}
                saved_docs_path.write_text(
                    json.dumps(saved_docs_serializable, ensure_ascii=False, indent=2),
                    encoding="utf-8"
                )
                
                checklist_path = job_dir / "results" / safe_name / "07_Dynamic_Checklist.md"
                run_cmd([
                    "python3", str(SCRIPTS_DIR / "build_dynamic_checklist.py"),
                    str(tender_ai_path),
                    str(forms_data_path),
                    str(checklist_path),
                    str(saved_docs_path)
                ])
                print(f"✓ Dynamic checklist built: {checklist_path.name}")
            except Exception as cl_err:
                print(f"⚠ Dynamic checklist generation failed: {cl_err}")
            
            update_job(job_id, stage="توليد المخططات والمناظير المعمارية (4 صور AI)...", progress=80)
            # ===== Architectural renders (LAST - most prone to failure) =====
            plans_dir = job_dir / "results" / safe_name / "03_Architectural"
            plans_dir.mkdir(parents=True, exist_ok=True)
            
            # Free up memory before image generation (most memory-intensive step)
            import gc
            gc.collect()
            
            try:
                generate_architectural(meta, plans_dir, safe_name)
            except MemoryError:
                print(f"⚠ OUT OF MEMORY during arch gen - using placeholders")
                gc.collect()
                try:
                    create_placeholder_images(plans_dir, meta)
                except Exception:
                    pass
            except Exception as e:
                print(f"Architectural generation skipped: {e}")
                try:
                    create_placeholder_images(plans_dir, meta)
                except Exception:
                    pass
            
            # README
            readme_path = job_dir / "results" / safe_name / "00_README.md"
            readme_path.write_text(build_tender_readme(meta, bidder_data), encoding="utf-8")
            
            all_outputs.append(safe_name)
        
        # Final ZIP
        final_zip = OUTPUTS_DIR / f"Musanada_Package_{job_id}.zip"
        with zipfile.ZipFile(final_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
            results_root = job_dir / "results"
            for fp in results_root.rglob("*"):
                if fp.is_file():
                    zf.write(fp, fp.relative_to(results_root))
        
        # Build response file list
        update_job(job_id, stage="تجميع الحزمة النهائية...", progress=95)
        
        files = [{
            "name": f"Musanada_Package_{job_id}.zip",
            "url": f"/outputs/Musanada_Package_{job_id}.zip",
            "description": f"الحزمة الكاملة ({len(all_outputs)} مناقصة)"
        }]
        
        for tender_name in all_outputs[:1]:
            results_dir = job_dir / "results" / tender_name
            for fp in (results_dir).rglob("*"):
                if fp.is_file() and fp.suffix.lower() in [".docx", ".xlsx", ".pdf", ".png", ".md"]:
                    rel = fp.relative_to(results_dir)
                    files.append({
                        "name": fp.name,
                        "url": f"/outputs/{job_id}/results/{tender_name}/{rel}",
                        "description": describe_file(fp.name)
                    })
        
        update_job(job_id,
            status="ok",
            stage="اكتمل!",
            progress=100,
            tenders_processed=len(all_outputs),
            files=files,
            finished_at=datetime.now().isoformat(),
        )
        print(f"✓✓✓ Job {job_id} completed: {len(all_outputs)} tenders, {len(files)} files")
        return  # background thread - no HTTP response
    
    except subprocess.CalledProcessError as e:
        traceback.print_exc()
        update_job(job_id,
            status="error",
            stage="خطأ",
            message=f"خطأ في تشغيل سكريبت: {Path(e.cmd[1]).name if len(e.cmd) > 1 else e.cmd[0]}",
            details=(e.stderr or "")[:500],
            finished_at=datetime.now().isoformat(),
        )
    except Exception as e:
        traceback.print_exc()
        update_job(job_id,
            status="error",
            stage="خطأ",
            message=str(e),
            error=traceback.format_exc()[:1000],
            finished_at=datetime.now().isoformat(),
        )


# ============ HELPERS ============
def run_cmd(cmd, cwd=None):
    print(f">>> Running: {Path(cmd[1]).name if len(cmd)>1 else ' '.join(cmd)}")
    result = subprocess.run(cmd, cwd=cwd, capture_output=True, text=True, timeout=300)
    if result.returncode != 0:
        print(f"STDERR: {result.stderr[:500]}")
        raise subprocess.CalledProcessError(result.returncode, cmd, result.stdout, result.stderr)
    return result


def describe_file(filename):
    lower = filename.lower()
    if "form_a" in lower: return "خطاب الغلاف"
    if "form_d" in lower: return "بيانات المزايد"
    if "form_e" in lower: return "الإعلان المحلف"
    if "form_g" in lower: return "خبرة المشاريع التفصيلية"
    if "form_h" in lower: return "خبرة المعماري"
    if "folder1" in lower: return "قائمة تدقيق - المتطلبات المسبقة"
    if "folder2" in lower: return "قائمة تدقيق - العرض الفني"
    if "folder3" in lower: return "قائمة تدقيق - العرض المالي"
    if "financial_model" in lower or ".xlsx" in lower: return "النموذج المالي"
    if "site_plan" in lower: return "المخطط الموقعي"
    if "perspective" in lower: return "منظور 3D"
    if "facade" in lower: return "الواجهة المعمارية"
    if "interior" in lower: return "المنظور الداخلي"
    if "analysis" in lower: return "التحليل العربي"
    if "readme" in lower: return "دليل الحزمة"
    return ""


def build_tender_readme(meta, bidder):
    return f"""# {meta.get('auction_title', 'Tender Package')}

> **حزمة تجهيز المناقصة الكاملة**  
> تم إعدادها بواسطة: مساندة للاستشارات الهندسية  
> التاريخ: {datetime.now().strftime('%Y-%m-%d %H:%M')}

---

## 🏢 الشركة المتقدمة
- **الاسم:** {bidder.get('company_legal_name', '-')}
- **الترخيص:** {bidder.get('trade_license_no', '-')}
- **الممثل المفوض:** {bidder.get('authorized_signatory_name', '-')} ({bidder.get('authorized_signatory_title', '-')})

## 📋 المناقصة
- **الجهة:** {meta.get('issuing_authority', '-')}
- **الموقع:** {meta.get('location', '-')}
- **المساحة:** {meta.get('land_area_sqm', '-')} م²
- **الحد الأدنى للإيجار:** AED {meta.get('min_rent_per_sqm', '-')}/م²
- **بدل المناقصة (+20%):** AED {float(meta.get('min_rent_per_sqm', 100)) * 1.2:.2f}/م²
- **آخر موعد:** {meta.get('closing_date', '-')}

## 📁 محتويات المجلد

| # | المجلد | الوصف |
|---|--------|------|
| 01 | `01_Forms/` | 8 نماذج جاهزة للتوقيع |
| 02 | `02_Financial_Model/` | نموذج Excel مالي بـ 10 شيتات |
| 03 | `03_Architectural/` | المخطط الموقعي + المنظور 3D |
| 04 | `04_Analysis/` | التحليل العربي + بيانات الإدخال |
"""


def generate_architectural(meta, output_dir, name):
    """Generate architectural visuals.
    
    Tries OpenAI DALL-E 3 first (production-friendly), falls back to placeholders.
    NEVER includes Arabic text in prompts to prevent broken glyphs.
    """
    facility_type = meta.get("facility_type", "Commercial Center")
    land_area = float(meta.get("land_area_sqm", 5000) or 5000)
    
    NO_TEXT = (
        " STRICTLY NO TEXT of any language. NO ARABIC, NO ENGLISH. "
        "NO LABELS, NO SIGNAGE WITH WRITING, NO LOGOS WITH TEXT. "
        "Pure visual architecture only — clean blank surfaces where signs would normally be."
    )
    
    prompts = {
        "01_Site_Plan.png": (
            f"Architectural top-down site plan of a {facility_type} facility, "
            f"plot area {land_area:.0f} square meters. "
            f"Clean technical CAD drawing on white background with navy blue lines. "
            f"Building footprints as geometric rectangles, parking lot with parking-space lines, "
            f"palm tree symbols, internal roads, pedestrian walkways, landscape zones. "
            f"Professional architectural site plan. No annotations, no labels, just shapes." + NO_TEXT,
            "1:1"
        ),
        "02_Perspective_3D.png": (
            f"Photorealistic 3D bird's-eye perspective rendering of a modern {facility_type} complex. "
            f"Contemporary minimalist Middle Eastern architecture. Beige stone and white facade. "
            f"Subtle gold metallic accents on edges. Large clear glass windows. "
            f"Manicured landscaping with palm trees, luxury cars in parking lot, "
            f"clear blue sky, golden hour sunlight casting long warm shadows. "
            f"Ultra-detailed award-winning architectural visualization. "
            f"Completely blank facade panels — no signage, no logos." + NO_TEXT,
            "16:9"
        ),
        "03_Facade_View.png": (
            f"Front facade eye-level photographic view of a modern {facility_type} building. "
            f"Contemporary architecture with clean geometric lines. "
            f"Beige stone exterior with large glass entrance and gold metal trim. "
            f"Landscaped entrance with palm trees, ambient evening LED lighting. "
            f"Professional architectural photography style. Completely blank facade — "
            f"no text, no signs, no logos visible anywhere." + NO_TEXT,
            "16:9"
        ),
        "04_Interior_Concept.png": (
            f"Interior reception lounge of a modern {facility_type}. "
            f"Spacious double-height ceiling, warm wood and stone finishes, "
            f"subtle gold metal pendant lighting, glass partitions, polished concrete floor, "
            f"comfortable lounge seating, large indoor plants. Premium minimalist design. "
            f"Architectural interior photography. Blank walls — no text artwork, "
            f"no signage." + NO_TEXT,
            "16:9"
        ),
    }
    
    import gc
    for filename, (prompt, aspect) in prompts.items():
        out_path = output_dir / filename
        try:
            success = generate_image_dalle(prompt, out_path)
        except Exception as e:
            print(f"⚠ Image gen failed for {filename}: {e}")
            success = False
        
        if not success:
            try:
                create_placeholder_image(out_path, filename)
            except Exception as e:
                print(f"⚠ Placeholder creation failed for {filename}: {e}")
        
        # Free memory between images (critical for 512MB Free tier)
        gc.collect()
        
        # Compress the image to save disk space
        try:
            compress_image(out_path, quality=80, max_dim=1280)
        except Exception:
            pass


def generate_image_dalle(prompt, output_path):
    """Generate via OpenAI gpt-image-2 (memory-efficient: streams to disk).
    
    Uses streaming + smallest possible size to fit in 512MB Free tier RAM.
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return False
    
    import gc
    
    try:
        import requests
        import base64
        
        # Use smaller sizes to keep memory low. 1024x1024 = ~1.5MB instead of 4MB
        for model, size in [
            ("gpt-image-2", "1024x1024"),     # Smallest gpt-image-2 size
            ("gpt-image-1", "1024x1024"),
            ("dall-e-3", "1024x1024"),
        ]:
            payload = {
                "model": model,
                "prompt": prompt[:4000],
                "n": 1,
                "size": size,
            }
            if model == "dall-e-3":
                payload["quality"] = "standard"
            
            try:
                # Use streaming to reduce peak memory
                with requests.post(
                    "https://api.openai.com/v1/images/generations",
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    },
                    json=payload,
                    timeout=180,
                    stream=True
                ) as response:
                    
                    if response.status_code != 200:
                        try:
                            err = response.json().get("error", {})
                            err_msg = err.get("message", "")
                        except:
                            err_msg = response.text[:200]
                        if "does not exist" in err_msg or "not found" in err_msg.lower():
                            print(f"  {model}: not available, trying next...")
                            continue
                        else:
                            print(f"  {model}: {err_msg[:100]}")
                            return False
                    
                    # Read response (small JSON wrapper)
                    data = response.json()
                
                item = data["data"][0]
                
                if "b64_json" in item and item["b64_json"]:
                    b64_str = item["b64_json"]
                    del data, item  # free response data immediately
                    
                    # Decode + write in one shot, then drop the strings
                    img_bytes = base64.b64decode(b64_str)
                    del b64_str
                    gc.collect()
                    
                    with open(output_path, "wb") as f:
                        f.write(img_bytes)
                    del img_bytes
                    gc.collect()
                    
                    size_kb = Path(output_path).stat().st_size // 1024
                    print(f"✓ Generated with {model} ({size_kb} KB)")
                    return True
                
                elif "url" in item and item["url"]:
                    url = item["url"]
                    del data
                    # Stream download to file (low memory)
                    with requests.get(url, timeout=60, stream=True) as img_response:
                        if img_response.status_code == 200:
                            with open(output_path, "wb") as f:
                                for chunk in img_response.iter_content(chunk_size=64 * 1024):
                                    f.write(chunk)
                            gc.collect()
                            print(f"✓ Generated with {model} via URL ({Path(output_path).stat().st_size // 1024} KB)")
                            return True
            
            except Exception as e:
                print(f"  {model} error: {e}")
                continue
        
        return False
    except Exception as e:
        print(f"Image gen failed: {e}")
        return False


def compress_image(image_path: Path, quality: int = 80, max_dim: int = 1280):
    """Compress image in place to reduce size. Saves a lot of memory + disk."""
    try:
        from PIL import Image
        with Image.open(image_path) as img:
            # Resize if too large
            if max(img.size) > max_dim:
                img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)
            
            # Save as JPEG-quality PNG (or convert to JPEG)
            if image_path.suffix.lower() == '.png':
                # Save as optimized PNG
                img.save(image_path, "PNG", optimize=True)
            else:
                img.save(image_path, quality=quality, optimize=True)
    except Exception as e:
        print(f"Could not compress {image_path.name}: {e}")


def create_placeholder_image(output_path, filename):
    """Create a professional placeholder when AI is unavailable."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        img = Image.new('RGB', (1600, 900), (245, 245, 240))
        d = ImageDraw.Draw(img)
        d.rectangle([30, 30, 1569, 869], outline=(11, 61, 122), width=6)
        
        # Inner border
        d.rectangle([60, 60, 1539, 839], outline=(201, 162, 76), width=2)
        
        # Title
        try:
            title_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 56)
            sub_font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
        except:
            title_font = ImageFont.load_default()
            sub_font = ImageFont.load_default()
        
        title_map = {
            "01_Site_Plan.png": "ARCHITECTURAL SITE PLAN",
            "02_Perspective_3D.png": "3D PERSPECTIVE VIEW",
            "03_Facade_View.png": "FACADE ELEVATION",
            "04_Interior_Concept.png": "INTERIOR CONCEPT",
        }
        title = title_map.get(filename, "ARCHITECTURAL VISUAL")
        d.text((800, 380), title, anchor="mm", fill=(11, 61, 122), font=title_font)
        d.text((800, 470), "Reserved for AI-generated rendering", anchor="mm", fill=(120, 120, 120), font=sub_font)
        d.text((800, 510), "Configure OPENAI_API_KEY to enable", anchor="mm", fill=(120, 120, 120), font=sub_font)
        
        # Logo placement
        d.text((800, 700), "MUSANADA ENGINEERING CONSULTANCY", anchor="mm", fill=(201, 162, 76), font=sub_font)
        
        img.save(output_path)
    except Exception as e:
        print(f"Placeholder creation failed: {e}")


def create_placeholder_images(output_dir, meta):
    """Create all 4 placeholder images."""
    for fname in ["01_Site_Plan.png", "02_Perspective_3D.png", 
                  "03_Facade_View.png", "04_Interior_Concept.png"]:
        create_placeholder_image(output_dir / fname, fname)


# ============ DOCUMENT INTELLIGENCE ============
def extract_text_from_file(file_path: Path, max_chars: int = 50000) -> str:
    """Extract text from PDF/DOCX/Image. Returns concatenated text."""
    suffix = file_path.suffix.lower()
    text = ""
    try:
        if suffix == ".pdf":
            result = subprocess.run(
                ["pdftotext", "-layout", str(file_path), "-"],
                capture_output=True, text=True, timeout=60
            )
            text = result.stdout
        elif suffix == ".docx":
            from docx import Document
            doc = Document(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    paragraphs.append(" | ".join(c.text for c in row.cells))
            text = "\n".join(paragraphs)
        elif suffix in [".xlsx", ".xls"]:
            try:
                from openpyxl import load_workbook
                wb = load_workbook(str(file_path), data_only=True)
                rows = []
                for ws in wb.worksheets:
                    rows.append(f"== Sheet: {ws.title} ==")
                    for row in ws.iter_rows(values_only=True):
                        vals = [str(v) if v is not None else "" for v in row]
                        if any(vals):
                            rows.append(" | ".join(vals))
                text = "\n".join(rows)
            except Exception:
                pass
        elif suffix in [".jpg", ".jpeg", ".png"]:
            # Could OCR via OpenAI Vision; for now, return placeholder
            text = "[Image file - OCR not implemented for placeholders]"
    except Exception as e:
        print(f"Text extraction failed for {file_path.name}: {e}")
    
    return text[:max_chars] if text else ""


def call_openai_extract(text: str, schema_instructions: str) -> dict:
    """Call OpenAI to extract structured data from text. Returns dict or {}."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key or not text.strip():
        return {}
    
    try:
        import requests
        prompt = f"""Extract structured data from the following document text.
Return ONLY valid JSON matching this schema:
{schema_instructions}

If a field is not found, use empty string "" or empty array [].
Do not include any explanation, only the JSON object.

Document text:
\"\"\"
{text[:30000]}
\"\"\"

JSON output:"""
        
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "model": "gpt-4o-mini",
                "messages": [
                    {"role": "system", "content": "You are a data extraction expert. Return only valid JSON."},
                    {"role": "user", "content": prompt}
                ],
                "temperature": 0.1,
                "response_format": {"type": "json_object"},
            },
            timeout=60
        )
        
        if response.status_code == 200:
            data = response.json()
            content = data["choices"][0]["message"]["content"]
            return json.loads(content)
        else:
            print(f"OpenAI extract failed: {response.status_code} - {response.text[:200]}")
    except Exception as e:
        print(f"OpenAI extraction error: {e}")
    
    return {}


def extract_license_data(file_path: Path) -> dict:
    """Extract data from Trade License PDF/Image."""
    text = extract_text_from_file(file_path)
    if not text:
        return {"_source": file_path.name, "_status": "no_text"}
    
    schema = """{
    "license_no": "Economic License Number (e.g., CN-1234567)",
    "trade_name_en": "Company name in English",
    "trade_name_ar": "Company name in Arabic",
    "establishment_date": "YYYY-MM-DD",
    "expiry_date": "YYYY-MM-DD",
    "legal_form": "e.g. LLC, Sole Proprietorship",
    "activities": "Main licensed activities (comma-separated)",
    "owner_name": "Owner / Authorized representative",
    "partners_nationality": "e.g. UAE",
    "issuing_authority": "e.g. ADDED Abu Dhabi"
}"""
    
    extracted = call_openai_extract(text, schema)
    extracted["_source"] = file_path.name
    return extracted


def extract_profile_data(file_path: Path) -> dict:
    """Extract data from Company Profile PDF/DOCX."""
    text = extract_text_from_file(file_path)
    if not text:
        return {"_source": file_path.name, "_status": "no_text"}
    
    schema = """{
    "short_name": "Brand / short name",
    "vision": "Vision statement (1 sentence)",
    "mission": "Mission statement (1 sentence)",
    "founded_year": "Year of establishment",
    "team_size": "Number of employees if mentioned",
    "phone": "Main contact phone",
    "email": "Main contact email",
    "website": "Website URL",
    "address": "HQ address",
    "key_services": ["service 1", "service 2"],
    "certifications": ["ISO 9001", "etc"],
    "key_clients": ["client 1", "client 2"]
}"""
    
    extracted = call_openai_extract(text, schema)
    extracted["_source"] = file_path.name
    return extracted


def extract_projects_data(file_path: Path) -> list:
    """Extract past projects from a portfolio PDF for Form G."""
    text = extract_text_from_file(file_path)
    if not text:
        return []
    
    schema = """{
    "projects": [
        {
            "name": "Project name",
            "location": "City / area",
            "client": "Client name (if mentioned)",
            "scope": "Develop, manage, operate / etc",
            "amount": 0,
            "start": "MM-YYYY or YYYY",
            "end": "MM-YYYY or YYYY",
            "gfa": 0,
            "gla": 0,
            "status": "Built 100% / Under construction X%",
            "occupancy": "Occupancy rate if mentioned"
        }
    ]
}"""
    
    extracted = call_openai_extract(text, schema)
    projects = extracted.get("projects", [])
    
    # Normalize and add defaults
    normalized = []
    for p in projects[:15]:  # max 15 projects
        normalized.append({
            "name": p.get("name", ""),
            "location": p.get("location", ""),
            "scope": p.get("scope", "Develop, manage and operate"),
            "amount": int(p.get("amount", 0)) if isinstance(p.get("amount"), (int, float)) else 0,
            "start": p.get("start", ""),
            "end": p.get("end", ""),
            "gfa": int(p.get("gfa", 0)) if isinstance(p.get("gfa"), (int, float)) else 0,
            "gla": int(p.get("gla", 0)) if isinstance(p.get("gla"), (int, float)) else 0,
            "status": p.get("status", "Built 100%"),
            "dev_role": "YES", "leasing_role": "YES", "mgmt_role": "YES",
            "floor_eff": "92%",
            "occ_2022": p.get("occupancy", "85%"),
            "occ_2023": p.get("occupancy", "90%"),
        })
    return normalized


def extract_financials_data(file_path: Path) -> dict:
    """Extract key financial figures."""
    text = extract_text_from_file(file_path)
    if not text:
        return {"_source": file_path.name}
    
    schema = """{
    "auditor_name": "Audit firm name",
    "fiscal_year_end": "YYYY-MM-DD of latest year",
    "total_revenue_latest": 0,
    "total_assets_latest": 0,
    "net_profit_latest": 0,
    "shareholders_equity": 0,
    "years_covered": ["2022", "2023", "2024"],
    "currency": "AED"
}"""
    
    extracted = call_openai_extract(text, schema)
    extracted["_source"] = file_path.name
    return extracted


def extract_bank_data(file_path: Path) -> dict:
    """Extract summary from bank statement."""
    text = extract_text_from_file(file_path)
    if not text:
        return {"_source": file_path.name}
    
    schema = """{
    "bank_name": "Bank name",
    "account_holder": "Account holder name",
    "account_number": "Account number (last 4 digits only)",
    "statement_period_start": "YYYY-MM-DD",
    "statement_period_end": "YYYY-MM-DD",
    "average_balance": 0,
    "currency": "AED"
}"""
    
    extracted = call_openai_extract(text, schema)
    extracted["_source"] = file_path.name
    return extracted


def build_attachments_index(output_path: Path, saved_docs: dict, intel: dict):
    """Build an index of all attached documents for the bidder."""
    lines = [
        "# 📎 فهرس المرفقات",
        "",
        "> جدول بكل الملفات اللي رفعتها وأماكنها في الحزمة + البيانات اللي استخرجها الـAI",
        "",
        "---",
        "",
        "## 📋 المرفقات حسب الفئة",
        "",
    ]
    
    category_info = {
        "license": ("📜 الرخصة التجارية", "01_Forms/Attachments_Technical/N8_Trade_License/"),
        "profile": ("📑 بروفايل الشركة", "01_Forms/Attachments_Technical/Bonus_Company_Profile/"),
        "works": ("🏗️ الأعمال السابقة", "01_Forms/Attachments_Technical/N18_Bidders_Experience/"),
        "id": ("🪪 هوية الممثل", "01_Forms/Attachments_Technical/N9_Power_of_Attorney/"),
        "orgchart": ("🗂️ الهيكل التنظيمي", "01_Forms/Attachments_Technical/N17_Organization_Chart/"),
        "financials": ("📊 الميزانية المدققة", "02_Financial_Model/Attachments_Financial/Audited_Financial_Statements/"),
        "bank": ("🏦 كشف بنكي", "02_Financial_Model/Attachments_Financial/Bank_Statements/"),
        "other": ("📎 مستندات أخرى", "05_Additional_Documents/"),
    }
    
    for cat, (label, location) in category_info.items():
        files = saved_docs.get(cat, [])
        if not files:
            continue
        lines.append(f"### {label}")
        lines.append(f"**الموقع:** `{location}`")
        lines.append("")
        for f in files:
            lines.append(f"- `{f.name}` ({f.stat().st_size // 1024} KB)")
        lines.append("")
    
    if intel:
        lines.append("---")
        lines.append("")
        lines.append("## 🤖 البيانات المستخرجة بواسطة AI")
        lines.append("")
        lines.append("هذه البيانات استخرجها الـAI تلقائياً من الملفات اللي رفعتها وتم استخدامها لملء الفورمات:")
        lines.append("")
        
        if "license" in intel:
            lines.append("### من الرخصة التجارية:")
            lic = intel["license"]
            for key, val in lic.items():
                if not key.startswith("_") and val:
                    lines.append(f"- **{key}:** {val}")
            lines.append("")
        
        if "profile" in intel:
            lines.append("### من بروفايل الشركة:")
            prof = intel["profile"]
            for key, val in prof.items():
                if not key.startswith("_") and val:
                    if isinstance(val, list):
                        lines.append(f"- **{key}:** {', '.join(map(str, val))}")
                    else:
                        lines.append(f"- **{key}:** {val}")
            lines.append("")
        
        if "projects" in intel and intel["projects"]:
            lines.append(f"### مشاريع مستخرجة من ملف الأعمال السابقة ({len(intel['projects'])} مشروع):")
            for p in intel["projects"][:10]:
                lines.append(f"- **{p.get('name', '?')}** — {p.get('location', '')} — AED {p.get('amount', 0):,}")
            lines.append("")
        
        if "financials" in intel:
            lines.append("### من الميزانية المدققة:")
            fin = intel["financials"]
            for key, val in fin.items():
                if not key.startswith("_") and val:
                    lines.append(f"- **{key}:** {val}")
            lines.append("")
        
        if "bank" in intel:
            lines.append("### من كشف الحساب البنكي:")
            bank = intel["bank"]
            for key, val in bank.items():
                if not key.startswith("_") and val:
                    lines.append(f"- **{key}:** {val}")
            lines.append("")
    
    lines.append("---")
    lines.append("")
    lines.append("> **ملاحظة:** راجع البيانات المستخرجة قبل التقديم — قد تحتاج تصحيح يدوي إذا كانت الملفات غير واضحة.")
    
    output_path.write_text("\n".join(lines), encoding="utf-8")


# ============ MAIN ============
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
